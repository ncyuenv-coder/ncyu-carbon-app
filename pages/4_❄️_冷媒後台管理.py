import streamlit as st
import pandas as pd
from datetime import datetime, date, timedelta
import gspread
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
import streamlit_authenticator as stauth
import plotly.express as px
import plotly.graph_objects as go
import time
import re
import io
import hashlib

try:
    from docx import Document
    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    import fitz  # PyMuPDF
    DOCX_READY = True
except ImportError:
    DOCX_READY = False

# ==========================================
# 0. 系統設定 (後台管理專用)
# ==========================================
st.set_page_config(page_title="冷媒後台管理 - 嘉義大學", page_icon="⚙️", layout="wide")

# ==========================================
# 1. CSS 樣式表
# ==========================================
st.markdown("""
<style>
    :root {
        color-scheme: light;
        --orange-bg: #E67E22;     
        --orange-dark: #D35400;
        --text-main: #2C3E50;
        --text-sub: #566573;
        --morandi-red: #C0392B; 
        --morandi-blue: #34495E;
    }
    [data-testid="stAppViewContainer"] { background-color: #EAEDED; color: var(--text-main); }
    [data-testid="stHeader"] { background-color: rgba(0,0,0,0); }
    [data-testid="stSidebar"] { background-color: #FFFFFF; border-right: 1px solid #BDC3C7; }
    
    div[data-baseweb="input"] > div, div[data-baseweb="base-input"] > input, textarea, input { background-color: #FFFFFF !important; border-color: #BDC3C7 !important; color: #000000 !important; font-size: 1.15rem !important; }
    div[data-baseweb="select"] > div { border-color: #BDC3C7 !important; background-color: #FFFFFF !important; }
    ul[data-baseweb="menu"] { background-color: #FFFFFF !important; }
    
    /* 主按鈕 (Primary) */
    button[kind="primary"] { background-color: var(--orange-bg) !important; color: #FFFFFF !important; border: 2px solid var(--orange-dark) !important; border-radius: 12px !important; font-size: 1.3rem !important; font-weight: 800 !important; padding: 0.7rem 1.5rem !important; box-shadow: 0 4px 6px rgba(230, 126, 34, 0.3) !important; width: 100%; }
    button[kind="primary"] p { color: #FFFFFF !important; } 
    button[kind="primary"]:hover { background-color: var(--orange-dark) !important; transform: translateY(-2px) !important; color: #FFFFFF !important; }
    
    /* 次按鈕 (Secondary) - 莫蘭迪深色底白字 */
    button[kind="secondary"] { background-color: #5D6D7E !important; color: #FFFFFF !important; border: 2px solid #34495E !important; border-radius: 12px !important; font-size: 1.15rem !important; font-weight: 800 !important; padding: 0.7rem 1.5rem !important; box-shadow: 0 4px 6px rgba(52, 73, 94, 0.3) !important; width: 100%; }
    button[kind="secondary"] p { color: #FFFFFF !important; } 
    button[kind="secondary"]:hover { background-color: #34495E !important; transform: translateY(-2px) !important; color: #FFFFFF !important; }
    button[kind="secondary"]:hover p { color: #FFFFFF !important; }
    
    button[data-baseweb="tab"] div p { font-size: 1.3rem !important; font-weight: 900 !important; color: var(--text-sub); }
    button[data-baseweb="tab"][aria-selected="true"] div p { color: #E67E22 !important; border-bottom: 3px solid #E67E22; }
    
    [data-testid="stDataFrame"] { font-size: 1.25rem !important; }
    [data-testid="stDataFrame"] div { font-size: 1.25rem !important; }
    
    /* Admin KPI 卡片 */
    .admin-kpi-card { background-color: #FFFFFF; border: 1px solid #BDC3C7; border-radius: 12px; overflow: hidden; box-shadow: 0 4px 10px rgba(0,0,0,0.1); height: 100%; text-align: center; margin-bottom: 20px; }
    .admin-kpi-header { padding: 10px; font-size: 1.2rem; font-weight: bold; color: #2C3E50; border-bottom: 1px solid rgba(0,0,0,0.1); }
    .admin-kpi-body { padding: 20px; }
    .admin-kpi-value { font-size: 2.8rem; font-weight: 900; color: #2C3E50; margin-bottom: 5px; }
    
    /* 修改單位為深灰色 */
    .admin-kpi-unit { font-size: 1.1rem; color: #333333 !important; font-weight: 800; margin-left: 5px; }
    
    .dashboard-main-title { font-size: 1.8rem; font-weight: 900; text-align: center; color: #2C3E50; margin-bottom: 20px; background-color: #F8F9F9; padding: 10px; border-radius: 10px; border: 1px solid #BDC3C7; }
    
    .stRadio div[role="radiogroup"] label { background-color: #D6EAF8 !important; border: 1px solid #AED6F1 !important; border-radius: 8px !important; padding: 8px 15px !important; margin-right: 10px !important; }
    .stRadio div[role="radiogroup"] label p { font-size: 1.0rem !important; font-weight: 800 !important; color: #154360 !important; }
    .stRadio div[role="radiogroup"] label[data-checked="true"] { background-color: #1A5276 !important; border-color: #1A5276 !important; }
    .stRadio div[role="radiogroup"] label[data-checked="true"] p { color: #FFFFFF !important; }
</style>
""", unsafe_allow_html=True)

# 2. 身份驗證
def clean_secrets(obj):
    if isinstance(obj, dict) or "AttrDict" in str(type(obj)): return {k: clean_secrets(v) for k, v in obj.items()}
    elif isinstance(obj, list): return [clean_secrets(i) for i in obj]
    return obj

if st.session_state.get("authentication_status") is not True:
    st.warning("🔒 請先至首頁 (Hello) 登入系統")
    st.stop()

username = st.session_state.get("username")
name = st.session_state.get("name")

# 安全管控：限制只有 admin 能夠進入
if username != 'admin':
    st.error("🚫 權限不足：此頁面僅供系統管理員使用。")
    st.stop()

try:
    _raw_creds = st.secrets["credentials"]
    credentials_login = clean_secrets(_raw_creds)
    cookie_cfg = st.secrets["cookie"]
    authenticator = stauth.Authenticate(credentials_login, cookie_cfg["name"], cookie_cfg["key"], cookie_cfg["expiry_days"])
except: pass

with st.sidebar:
    st.header(f"👤 {name} (管理員)")
    st.success("☁️ 後台連線正常")
    authenticator.logout('登出系統', 'sidebar')

# 3. 資料庫連線
REF_SHEET_ID = "1p7GsW-nrjerXhnn3pNgZzu_CdIh1Yxsm-fLJDqQ6MqA"

@st.cache_resource
def init_google_ref():
    oauth = st.secrets["gcp_oauth"]
    creds = Credentials(token=None, refresh_token=oauth["refresh_token"], token_uri="https://oauth2.googleapis.com/token", client_id=oauth["client_id"], client_secret=oauth["client_secret"], scopes=["https://www.googleapis.com/auth/drive", "https://www.googleapis.com/auth/spreadsheets"])
    gc = gspread.authorize(creds); drive = build('drive', 'v3', credentials=creds)
    return gc, drive

try:
    gc, drive_service = init_google_ref()
    sh_ref = gc.open_by_key(REF_SHEET_ID)
    try: ws_records = sh_ref.worksheet("冷媒填報紀錄")
    except: 
        ws_records = sh_ref.add_worksheet(title="冷媒填報紀錄", rows="1000", cols="15")
except Exception as e:
    st.error(f"❌ 資料庫連線失敗: {e}")
    st.stop()

DATA_GWP = {
    'HFC-1234yf (R-1234yf)': 0.0, 'HFC-125 (R-125)': 3170.0, 'HFC-134a (R-134a)': 1300.0,
    'HFC-143a (R-143a)': 4800.0, 'HFC-245fa (R-245fa)': 858.0, 'R404a': 3942.8,
    'R407c': 1624.21, 'R-407D': 1487.05, 'R408a': 2429.9, 'R410a': 1923.5,
    'R-507A': 3985.0, 'R-508A': 11607.0, 'R-508B': 11698.0, 'HFC-23 (R-23)': 12400.0,
    'HFC-32 (R-32)': 677.0, 'R-411A': 0.0, 'R-402A': 0.0, '其他': 0.0
}

def load_static_data_cloud():
    try:
        ws_coef = sh_ref.worksheet("冷媒係數表")
        df_coef = pd.DataFrame(ws_coef.get_all_records())
        gwp_map = {}
        if not df_coef.empty:
            for _, row in df_coef.iterrows():
                r_name_full = str(row.iloc[1]).strip().replace('\u3000', ' ').replace('\xa0', ' ')
                try: gwp = float(str(row.iloc[2]).replace(',', ''))
                except: gwp = 0.0
                if r_name_full:
                    gwp_map[r_name_full] = gwp
        if '其他' not in gwp_map:
            gwp_map['其他'] = 0.0
        return gwp_map
    except:
        return DATA_GWP

@st.cache_data(ttl=60)
def load_records_data():
    try:
        data = ws_records.get_all_values()
        if len(data) > 1:
            raw_headers = data[0]
            col_mapping = {}
            for h in raw_headers:
                clean_h = str(h).strip()
                if "填充量" in clean_h or "重量" in clean_h: col_mapping[h] = "冷媒填充量"
                elif "種類" in clean_h or "品項" in clean_h: col_mapping[h] = "冷媒種類"
                elif "日期" in clean_h or "維修" in clean_h: col_mapping[h] = "維修日期"
                else: col_mapping[h] = clean_h
            
            df = pd.DataFrame(data[1:], columns=raw_headers)
            df.rename(columns=col_mapping, inplace=True)
            return df
        else:
            return pd.DataFrame(columns=["填報時間","填報人","填報人分機","校區","所屬單位","填報單位名稱","建築物名稱","辦公室編號","維修日期","設備類型","設備品牌型號","冷媒種類","冷媒填充量","備註","佐證資料"])
    except Exception as e:
        return pd.DataFrame()

# ==========================================
# 4. 輔助函數：Google Drive 與 Word 匯出
# ==========================================
def get_drive_id(url):
    match = re.search(r'/d/([a-zA-Z0-9_-]+)', str(url))
    if match: return match.group(1)
    match = re.search(r'id=([a-zA-Z0-9_-]+)', str(url))
    if match: return match.group(1)
    return None

def download_and_convert_drive_files(drive_service_obj, file_id):
    try:
        request = drive_service_obj.files().get_media(fileId=file_id)
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while done is False:
            status, done = downloader.next_chunk()
        fh.seek(0)
        header = fh.read(4)
        fh.seek(0)
        if header == b'%PDF':
            images = []
            try:
                doc = fitz.open(stream=fh.read(), filetype="pdf")
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    pix = page.get_pixmap(dpi=150)
                    img_bytes = pix.tobytes("png")
                    images.append(io.BytesIO(img_bytes))
                return images
            except: return []
        return [fh]
    except: return []

def export_ref_docx(df_year, drive_srv):
    doc = Document()
    for section in doc.sections:
        section.page_width = Cm(21.0); section.page_height = Cm(29.7)
        section.left_margin = Cm(1.5); section.right_margin = Cm(1.5); section.top_margin = Cm(1.5); section.bottom_margin = Cm(1.5)

    df_export = df_year.copy()
    # 確保日期格式供排序
    df_export['維修日期'] = pd.to_datetime(df_export['維修日期']).dt.date
    # 依冷媒種類、發票日期排序
    df_export = df_export.sort_values(by=['冷媒種類', '維修日期'])

    global_seen_fids = set()
    global_image_hashes = set()

    for _, row in df_export.iterrows():
        p = doc.add_paragraph()
        p.add_run(f"所在校區：{row.get('校區', '-')} | 填報單位：{row.get('填報單位名稱', '-')}\n").bold = True
        p.add_run(f"冷媒種類：{row.get('冷媒種類', '-')} | 冷媒填充量：{row.get('冷媒填充量', 0)} 公斤\n").bold = True
        p.add_run(f"設備類型：{row.get('設備類型', '-')} | 建築物名稱：{row.get('建築物名稱', '-')} | 辦公室編號：{row.get('辦公室編號', '-')}\n").bold = True
        p.add_run(f"發票日期：{row.get('維修日期', '-')}\n").bold = True

        link_str = str(row.get('佐證資料', ''))
        images_to_print = []
        skipped_dup = False
        
        for link in link_str.split('\n'):
            link = link.strip()
            if not link or link in ["無", ""]: continue
            fid = get_drive_id(link)
            if fid:
                if fid not in global_seen_fids:
                    global_seen_fids.add(fid)
                    downloaded_imgs = download_and_convert_drive_files(drive_srv, fid)
                    for img_io in downloaded_imgs:
                        img_io.seek(0)
                        img_hash = hashlib.md5(img_io.read()).hexdigest()
                        img_io.seek(0)
                        if img_hash not in global_image_hashes:
                            global_image_hashes.add(img_hash)
                            images_to_print.append(img_io)
                        else:
                            skipped_dup = True
                else:
                    skipped_dup = True

        if len(images_to_print) == 1:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try: p_img.add_run().add_picture(images_to_print[0], width=Cm(16.0))
            except: pass
        elif len(images_to_print) > 1:
            table = doc.add_table(rows=0, cols=2)
            table.autofit = False
            for i, img in enumerate(images_to_print):
                if i % 2 == 0: row_cells = table.add_row().cells
                try:
                    p_img = row_cells[i % 2].paragraphs[0]
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.add_run().add_picture(img, width=Cm(8.0))
                except: pass
        
        if skipped_dup and not images_to_print:
            p_dup = doc.add_paragraph()
            p_dup.add_run("*(此佐證資料與前方共用或已顯示過，為節省篇幅自動省略)*").italic = True
        
        doc.add_page_break()
        
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

MORANDI_PALETTE = ['#88B04B', '#92A8D1', '#F7CAC9', '#5D6D7E', '#7FB3D5', '#E59866', '#F7DC6F', '#CCD1D1', '#76D7C4', '#F1948A', '#BB8FCE']

# ==========================================
# 5. 分頁 Fragment 模組 (防止跳轉)
# ==========================================
@st.fragment
def render_tab1_dashboard(df_clean_dash, all_years):
    st.markdown("<br>", unsafe_allow_html=True)
    c_year, _ = st.columns([1, 3])
    sel_year = c_year.selectbox("📅 選擇統計年份", all_years, key="dash_year")
    st.markdown("---")
    
    df_year = df_clean_dash[df_clean_dash['年份'] == sel_year] if not df_clean_dash.empty else pd.DataFrame()
    
    if not df_year.empty:
        st.markdown(f"<div class='dashboard-main-title'>{sel_year}年度 冷媒填充與碳排統計</div>", unsafe_allow_html=True)
        
        total_kg = df_year['冷媒填充量'].sum()
        total_co2_t = df_year['排放量(公噸)'].sum()
        count = len(df_year)
        
        k1, k2, k3 = st.columns(3)
        k1.markdown(f"""<div class="admin-kpi-card"><div class="admin-kpi-header" style="background-color: #A9CCE3;">📄 年度申報筆數</div><div class="admin-kpi-body"><div class="admin-kpi-value">{count}</div></div></div>""", unsafe_allow_html=True)
        k2.markdown(f"""<div class="admin-kpi-card"><div class="admin-kpi-header" style="background-color: #F5CBA7;">⚖️ 年度總填充量</div><div class="admin-kpi-body"><div class="admin-kpi-value">{total_kg:,.2f}<span class="admin-kpi-unit">公斤</span></div></div></div>""", unsafe_allow_html=True)
        k3.markdown(f"""<div class="admin-kpi-card"><div class="admin-kpi-header" style="background-color: #E6B0AA;">☁️ 年度總碳排放量</div><div class="admin-kpi-body"><div class="admin-kpi-value">{total_co2_t:,.4f}<span class="admin-kpi-unit">公噸CO<sub>2</sub>e</span></div></div></div>""", unsafe_allow_html=True)
        
        st.markdown("---")
        
        st.subheader("📈 年度冷媒填充概況")
        campus_opts = ["全校", "蘭潭校區", "民雄校區", "新民校區", "林森校區"]
        f_campus_1 = st.radio("填充概況校區選擇", campus_opts, horizontal=True, key="radio_c1", label_visibility="collapsed")
        
        df_c1 = df_year.copy()
        if f_campus_1 != "全校":
            df_c1 = df_c1[df_c1['校區'] == f_campus_1]
        
        if not df_c1.empty:
            c1_group = df_c1.groupby(['冷媒顯示名稱', '設備類型'])['冷媒填充量'].sum().reset_index()
            c1_counts = c1_group.groupby('冷媒顯示名稱')['設備類型'].count() 
            c1_totals = c1_group.groupby('冷媒顯示名稱')['冷媒填充量'].sum().reset_index() 
            c1_group['label_text'] = c1_group.apply(lambda row: f"{row['冷媒填充量']:.2f}" if c1_counts[row['冷媒顯示名稱']] > 1 else None, axis=1)

            fig1 = px.bar(c1_group, x='冷媒顯示名稱', y='冷媒填充量', color='設備類型', 
                          text='label_text', color_discrete_sequence=MORANDI_PALETTE)
            
            fig1.add_trace(go.Scatter(
                x=c1_totals['冷媒顯示名稱'], 
                y=c1_totals['冷媒填充量'],
                text=c1_totals['冷媒填充量'].apply(lambda x: f'{x:.2f}'),
                mode='text',
                textposition='top center',
                textfont=dict(size=18, color='black'),
                showlegend=False,
                hoverinfo='skip'
            ))

            fig1.update_layout(height=600, yaxis_title="冷媒填充量(公斤)", xaxis_title="冷媒種類", font=dict(size=20), showlegend=True, margin=dict(t=50))
            fig1.update_xaxes(tickfont=dict(size=18, color='#000000'))
            fig1.update_yaxes(tickfont=dict(size=18, color='#000000'))
            fig1.update_traces(selector=dict(type='bar'), width=0.5, textfont_size=16, textposition='inside', textangle=0)
            st.plotly_chart(fig1, use_container_width=True)
        else:
            st.info("無資料")

        st.markdown("---")
        
        st.subheader("🏆 年度前十大填充單位")
        top_units = df_year.groupby('填報單位名稱')['冷媒填充量'].sum().nlargest(10).index.tolist()
        df_top10 = df_year[df_year['填報單位名稱'].isin(top_units)].copy()
        
        if not df_top10.empty:
            c2_group = df_top10.groupby(['填報單位名稱', '冷媒顯示名稱'])['冷媒填充量'].sum().reset_index()
            c2_counts = c2_group.groupby('填報單位名稱')['冷媒顯示名稱'].count()
            c2_totals = c2_group.groupby('填報單位名稱')['冷媒填充量'].sum().reset_index()
            c2_group['label_text'] = c2_group.apply(lambda row: f"{row['冷媒填充量']:.2f}" if c2_counts[row['填報單位名稱']] > 1 else None, axis=1)
            
            fig2 = px.bar(c2_group, x='填報單位名稱', y='冷媒填充量', color='冷媒顯示名稱',
                          text='label_text', color_discrete_sequence=MORANDI_PALETTE)
            
            fig2.add_trace(go.Scatter(
                x=c2_totals['填報單位名稱'], 
                y=c2_totals['冷媒填充量'],
                text=c2_totals['冷媒填充量'].apply(lambda x: f'{x:.2f}'),
                mode='text',
                textposition='top center',
                textfont=dict(size=18, color='black'),
                showlegend=False,
                hoverinfo='skip'
            ))

            fig2.update_layout(height=600, xaxis={'categoryorder':'total descending'}, yaxis_title="冷媒填充量(公斤)", font=dict(size=18), margin=dict(t=50))
            fig2.update_xaxes(tickfont=dict(size=16, color='#000000'))
            fig2.update_yaxes(tickfont=dict(size=16, color='#000000'))
            fig2.update_traces(selector=dict(type='bar'), width=0.5, textfont_size=14, textposition='inside', textangle=0)
            st.plotly_chart(fig2, use_container_width=True)
        else:
            st.info("無資料")
        
        st.markdown("---")
        
        st.subheader("🍩 冷媒填充資訊分析")
        f_campus_3 = st.radio("資訊分析校區選擇", campus_opts, horizontal=True, key="radio_c3", label_visibility="collapsed")
        df_c3 = df_year.copy()
        if f_campus_3 != "全校":
            df_c3 = df_c3[df_c3['校區'] == f_campus_3]
        
        if not df_c3.empty:
            st.markdown("##### 1. 冷媒種類填充量佔比")
            # 依需求改為水平長條圖，顯示 項目、填充量、百分比
            type_kg = df_c3.groupby('冷媒顯示名稱')['冷媒填充量'].sum().reset_index()
            type_kg = type_kg[type_kg['冷媒填充量'] > 0].sort_values('冷媒填充量', ascending=True)
            if not type_kg.empty:
                tot_kg = type_kg['冷媒填充量'].sum()
                type_kg['Label'] = type_kg.apply(lambda x: f"{x['冷媒顯示名稱']}: {x['冷媒填充量']:.2f} ({x['冷媒填充量']/tot_kg*100:.1f}%)", axis=1)
                
                fig3a = px.bar(type_kg, x='冷媒填充量', y='冷媒顯示名稱', orientation='h', text='Label', color='冷媒顯示名稱', color_discrete_sequence=MORANDI_PALETTE)
                fig3a.update_layout(height=500, showlegend=False, xaxis_title="填充量 (公斤)", yaxis_title="", font=dict(size=16), plot_bgcolor='rgba(0,0,0,0)')
                fig3a.update_xaxes(showgrid=True, gridcolor='#EAEDED', tickfont=dict(size=16, color='#333333'), title_font=dict(size=18, color='#333333'))
                fig3a.update_yaxes(tickfont=dict(size=16, color='#333333'))
                fig3a.update_traces(textposition='outside', textfont=dict(size=18, color='black'))
                st.plotly_chart(fig3a, use_container_width=True)
            else:
                st.info("無冷媒種類資料")
            
            st.markdown("<br>", unsafe_allow_html=True)
            
            st.markdown("##### 2. 冷媒填充設備類型統計")
            c3_l, c3_r = st.columns(2)
            
            eq_count = df_c3.groupby('設備類型')['冷媒填充量'].count().reset_index(name='count')
            fig3b_l = px.pie(eq_count, values='count', names='設備類型', title='依填充次數統計', hole=0.4,
                             color_discrete_sequence=MORANDI_PALETTE)
            fig3b_l.update_layout(font=dict(size=18), legend=dict(font=dict(size=16)), uniformtext_minsize=16, uniformtext_mode='show')
            fig3b_l.update_traces(textinfo='label+percent', textfont_size=18, textposition='outside',
                                  hovertemplate='<b>%{label}</b><br>填充次數: %{value} 次<br>佔比: %{percent:.1%}<extra></extra>')
            
            eq_weight = df_c3.groupby('設備類型')['冷媒填充量'].sum().reset_index()
            fig3b_r = px.pie(eq_weight, values='冷媒填充量', names='設備類型', title='依填充重量統計', hole=0.4,
                             color_discrete_sequence=MORANDI_PALETTE)
            fig3b_r.update_layout(font=dict(size=18), legend=dict(font=dict(size=16)), uniformtext_minsize=16, uniformtext_mode='show')
            fig3b_r.update_traces(textinfo='label+percent', textfont_size=18, textposition='outside',
                                  hovertemplate='<b>%{label}</b><br>填充重量: %{value:.2f} kg<br>佔比: %{percent:.1%}<extra></extra>')
            
            with c3_l: st.plotly_chart(fig3b_l, use_container_width=True)
            with c3_r: st.plotly_chart(fig3b_r, use_container_width=True)
        else:
            st.info("無資料")
        
        st.markdown("---")
        
        st.subheader("🌍 全校冷媒填充碳排放量(公噸二氧化碳當量)結構")
        fig_tree = px.treemap(df_year, path=['校區', '填報單位名稱'], values='排放量(公噸)', 
                              color='校區', color_discrete_sequence=MORANDI_PALETTE)
        fig_tree.update_traces(texttemplate='%{label}<br>%{value:.4f}<br>%{percentRoot:.1%}', textfont=dict(size=24))
        st.plotly_chart(fig_tree, use_container_width=True)
        
    else:
        st.info("該年度無資料")

@st.fragment
def render_tab2_edit(df_clean, df_records):
    st.markdown("<br>", unsafe_allow_html=True)
    st.subheader("📝 申報資料異動")
    
    if st.button("🔄 更新背景資料庫 (從 Google Sheet 同步)", key="btn_update_db"):
        with st.spinner("正在從雲端下載最新資料..."):
            st.session_state['gwp_map'] = load_static_data_cloud()
        st.success("✅ 資料庫已更新！")

    if not df_clean.empty:
        edited = st.data_editor(
            df_clean, 
            column_config={
                "佐證資料": st.column_config.LinkColumn("佐證"),
                "冷媒填充量": st.column_config.NumberColumn("填充量", min_value=0.0, step=0.1),
                "維修日期": st.column_config.DateColumn("日期", format="YYYY-MM-DD")
            },
            num_rows="dynamic",
            use_container_width=True,
            key="ref_editor"
        )
        
        if st.button("💾 儲存變更", type="primary"):
            try:
                df_final = edited.copy()
                cols_to_remove = ['年份', '月份', '排放量(kgCO2e)', '排放量(公噸)', '冷媒顯示名稱']
                for c in cols_to_remove:
                    if c in df_final.columns: del df_final[c]
                
                df_final['維修日期'] = df_final['維修日期'].astype(str)
                cols_to_write = df_records.columns.tolist()
                cols_to_write = [c for c in cols_to_write if c in df_final.columns]
                df_final = df_final[cols_to_write]
                
                ws_records.clear()
                ws_records.update([df_final.columns.tolist()] + df_final.astype(str).values.tolist())
                st.success("✅ 資料更新成功！")
                st.cache_data.clear()
            except Exception as e:
                st.error(f"更新失敗: {e}")
    else:
        st.info("無資料可編輯")

@st.fragment
def render_tab3_export(df_clean_dash, all_years):
    st.markdown("<br>", unsafe_allow_html=True)
    st.subheader("📁 年度冷媒填充統計及佐證下載")
    
    c_year, _ = st.columns([1, 3])
    sel_admin_year = c_year.selectbox("📅 請選擇下載年度", all_years, index=0, key="export_year")
    st.markdown("---")

    df_year = df_clean_dash[df_clean_dash['年份'] == sel_admin_year] if not df_clean_dash.empty else pd.DataFrame()
    
    if not df_year.empty:
        st.markdown("##### 📊 選擇您要下載的資料類型")
        if not DOCX_READY:
            st.error("⚠️ 系統尚未安裝 `python-docx` 或 `pymupdf` 套件。請確認已將其加入 `requirements.txt` 並重新部署。")
        else:
            c1, c2 = st.columns(2)
            with c1:
                # 下載 CSV
                cols_to_keep = [c for c in df_year.columns if c not in ['年份', '月份', '冷媒顯示名稱']]
                df_csv = df_year[cols_to_keep]
                csv_data = df_csv.to_csv(index=False).encode('utf-8-sig')
                st.download_button(
                    label=f"📥 年度冷媒填充統計表 (CSV)", 
                    data=csv_data, 
                    file_name=f"{sel_admin_year}_冷媒填充統計表.csv", 
                    mime="text/csv", 
                    use_container_width=True
                )
            with c2:
                # 下載 Word
                if st.button("⚡ 產生冷媒填充佐證 (Word)", use_container_width=True):
                    with st.spinner("正在下載圖片並合併 (自動進行 MD5 圖像去重)，可能需要幾分鐘..."):
                        st.session_state['doc_ref'] = export_ref_docx(df_year, drive_service)
                if 'doc_ref' in st.session_state:
                    st.download_button(
                        "⬇️ 下載冷媒填充佐證", 
                        data=st.session_state['doc_ref'], 
                        file_name=f"{sel_admin_year}_冷媒填充佐證資料.docx", 
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", 
                        use_container_width=True
                    )
    else:
        st.info(f"{sel_admin_year} 年度尚無資料。")


# ==========================================
# 6. 主程式 (管理員後台)
# ==========================================
def render_admin_dashboard():
    st.markdown("### 👑 冷媒管理後台")
    
    if 'gwp_map' not in st.session_state:
        st.session_state['gwp_map'] = DATA_GWP
        
    gwp_map = st.session_state['gwp_map']
    df_records = load_records_data()

    admin_tabs = st.tabs(["📊 全校冷媒填充儀表板", "📝 申報資料異動", "📁 年度冷媒填充統計及佐證下載"])

    # 預處理資料 (供儀表板與下載使用)
    df_clean_dash = df_records.copy()
    if not df_clean_dash.empty:
        campus_map = {"林森校區-民國路": "林森校區", "社口林場": "蘭潭校區"}
        df_clean_dash['校區'] = df_clean_dash['校區'].replace(campus_map)
        df_clean_dash['冷媒顯示名稱'] = df_clean_dash['冷媒種類']
        df_clean_dash['冷媒填充量'] = pd.to_numeric(df_clean_dash['冷媒填充量'], errors='coerce').fillna(0)
        df_clean_dash['維修日期'] = pd.to_datetime(df_clean_dash['維修日期'], errors='coerce')
        df_clean_dash['年份'] = df_clean_dash['維修日期'].dt.year.fillna(datetime.now().year).astype(int)
        df_clean_dash['月份'] = df_clean_dash['維修日期'].dt.month.fillna(0).astype(int)
        df_clean_dash['排放量(kgCO2e)'] = df_clean_dash.apply(lambda r: r['冷媒填充量'] * gwp_map.get(str(r['冷媒種類']).strip(), 0), axis=1)
        df_clean_dash['排放量(公噸)'] = df_clean_dash['排放量(kgCO2e)'] / 1000.0
        df_clean_dash['維修日期_Date'] = df_clean_dash['維修日期'].dt.date

    all_years = sorted(df_clean_dash['年份'].unique(), reverse=True) if not df_clean_dash.empty else [datetime.now().year]

    # 套用 fragment 防跳轉
    with admin_tabs[0]: 
        render_tab1_dashboard(df_clean_dash, all_years)
        
    with admin_tabs[1]: 
        df_clean_edit = df_clean_dash.copy()
        if not df_clean_edit.empty:
            df_clean_edit['維修日期'] = df_clean_edit['維修日期_Date']
            del df_clean_edit['維修日期_Date']
        render_tab2_edit(df_clean_edit, df_records)
        
    with admin_tabs[2]: 
        render_tab3_export(df_clean_dash, all_years)

if __name__ == "__main__":
    render_admin_dashboard()