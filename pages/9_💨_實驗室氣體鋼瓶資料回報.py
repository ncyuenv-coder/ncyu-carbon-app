import streamlit as st
import pandas as pd
import gspread
import datetime
import time
import io
import os
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn

# ================= 系統與全域參數設定 =================
st.set_page_config(page_title="實驗室氣體鋼瓶資料回報", page_icon="💨", layout="wide")

SHEET_ID = '1Hw4rXo4ww7O9YXTwoUJeWioO5ZzM_bivRcLLpOl26DY'
FOLDER_ID = '1cpj6l-zTgqLwphMFzLvvr53-Alf60idN'

CAMPUS_OPTS = ["蘭潭校區", "民雄校區", "新民校區", "林森校區"]
GAS_TYPES = ["請選擇", "二氧化碳", "甲烷", "乙炔", "一氧化二氮(笑氣)"]

# ================= 莫蘭迪風格 CSS & UI =================
def apply_morandi_theme():
    st.markdown("""
    <style>
        /* 字體全面放大一號 */
        label, p, .st-emotion-cache-1y4p8pa, .st-emotion-cache-1104ytp p, .stSelectionWithIcon p { font-size: 20px !important; }
        .stTextInput input, .stSelectbox div { font-size: 18px !important; }
        
        .title-box { background-color: #5C6B73; color: #FFFFFF; padding: 20px; border-radius: 10px; text-align: center; margin-bottom: 20px; }
        .section-header { background-color: #9DB4AB; color: #FFFFFF; padding: 12px 18px; border-radius: 8px; font-size: 22px; font-weight: bold; margin-top: 25px; margin-bottom: 15px; }
        .card { background-color: #FDFBF7; border: 1px solid #DCE4E3; padding: 20px; border-radius: 8px; margin-bottom: 15px; }
        
        /* 淺藍色 Radio 按鈕 (有無購買) */
        div.stRadio > div[role="radiogroup"] > label {
            background-color: #E3F2FD !important;
            padding: 12px 24px !important;
            border-radius: 8px !important;
            border: 2px solid #90CAF9 !important;
            cursor: pointer;
        }
        div.stRadio > div[role="radiogroup"] > label p { color: #0D47A1 !important; font-weight: bold !important; margin: 0;}
        div.stRadio > div[role="radiogroup"] > label[data-checked="true"] { background-color: #90CAF9 !important; border-color: #1565C0 !important; }
        
        /* 上傳檔案區域：全淺藍色底色 */
        [data-testid="stFileUploader"] { background-color: #E3F2FD !important; border: 2px dashed #64B5F6 !important; border-radius: 8px; padding: 20px; }
        [data-testid="stFileUploaderDropzone"] { background-color: transparent !important; border: none !important; }
        [data-testid="stFileUploaderDropzone"] > div { background-color: transparent !important; }
        
        /* 一般按鈕樣式 */
        div.stButton > button { background-color: #5C6B73 !important; color: #FFFFFF !important; font-weight: bold !important; font-size: 18px !important; padding: 12px 20px !important; border-radius: 6px !important;}
        div.stButton > button:hover { background-color: #4A565C !important; }
        
        /* 特製橘色下載按鈕 */
        div.stButton > button.orange-btn, div[data-testid="stDownloadButton"] button { 
            background-color: #E67E22 !important; border: 2px solid #D35400 !important; color: white !important; font-size: 18px !important; font-weight: bold !important; padding: 12px 20px !important; border-radius: 6px !important;
        }
        div.stButton > button.orange-btn:hover, div[data-testid="stDownloadButton"] button:hover { background-color: #D35400 !important; }
        
        [data-testid="stSidebar"], [data-testid="collapsedControl"], [data-testid="stHeader"] { display: none !important; }
        .block-container { padding-top: 2rem !important; }
    </style>
    """, unsafe_allow_html=True)

# ================= 🚀 PDF 轉圖片核心函式 =================
def get_pdf_preview_images(file_obj):
    """將 PDF 的「所有頁面」擷取為 PNG 圖片格式清單"""
    try:
        import fitz  # 需要 PyMuPDF 套件
        file_obj.seek(0)
        doc = fitz.open(stream=file_obj.read(), filetype="pdf")
        images = []
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            pix = page.get_pixmap(dpi=150) # 解析度 150 dpi
            images.append(io.BytesIO(pix.tobytes("png")))
        file_obj.seek(0) 
        return images
    except Exception as e:
        return []

# ================= Google API 與函式 =================
def with_retry(max_retries=3, base_delay=1.5):
    def decorator(func):
        def wrapper(*args, **kwargs):
            retries = 0
            while True:
                try: return func(*args, **kwargs)
                except Exception as e:
                    if retries >= max_retries: raise e
                    time.sleep(base_delay * (2 ** retries)); retries += 1
        return wrapper
    return decorator

@st.cache_resource
def get_credentials():
    oauth_secrets = st.secrets["gcp_oauth"]
    return Credentials(token=None, refresh_token=oauth_secrets["refresh_token"], client_id=oauth_secrets["client_id"], client_secret=oauth_secrets["client_secret"], token_uri="https://oauth2.googleapis.com/token")

def get_gc(): return gspread.authorize(get_credentials())

@st.cache_data(ttl=120, show_spinner=False)
def fetch_tracker_records(): 
    try: return get_gc().open_by_key(SHEET_ID).worksheet('發送紀錄與金鑰').get_all_records()
    except: return []

@st.cache_data(ttl=120, show_spinner=False)
def fetch_main_inventory(): 
    try: return get_gc().open_by_key(SHEET_ID).worksheet('氣體鋼瓶資料庫').get_all_records()
    except: return []

@with_retry(max_retries=3)
def append_dicts(worksheet, list_of_dicts):
    headers = worksheet.row_values(1)
    if headers and list_of_dicts: 
        worksheet.append_rows([[str(d.get(str(h).strip(), "")) for h in headers] for d in list_of_dicts])

@with_retry(max_retries=3)
def safe_delete_row(worksheet, idx): worksheet.delete_rows(idx)

@with_retry(max_retries=3)
def upload_to_drive(uploaded_file, file_name):
    creds = get_credentials(); drive_service = build('drive', 'v3', credentials=creds)
    file_metadata = {'name': file_name, 'parents': [FOLDER_ID]}
    uploaded_file.seek(0)
    media = MediaIoBaseUpload(io.BytesIO(uploaded_file.getvalue()), mimetype=uploaded_file.type, resumable=True)
    file = drive_service.files().create(body=file_metadata, media_body=media, fields='id').execute()
    return f"https://drive.google.com/file/d/{file.get('id')}/view?usp=sharing"

# ================= 🚀 產生 Word 申報表 (精準修復 AttributeError) =================
def create_report_docx(base_data, pur_data, status):
    doc = Document()
    
    # --- 安全設定字體的方法 ---
    def set_font_safe(run):
        run.font.name = 'Times New Roman'
        try:
            run.font._element.rFonts.set(qn('w:eastAsia'), '標楷體')
        except:
            pass

    # --- 全域字體設定 ---
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    try:
        style.font._element.rFonts.set(qn('w:eastAsia'), '標楷體')
    except:
        pass
    
    # --- 版面設定為「中等邊界」 ---
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    title = doc.add_heading('溫室氣體盤查 - 實驗室氣體鋼瓶申報表', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title.runs: set_font_safe(title.runs[0])
    
    doc.add_paragraph(f"盤查年度：2026年")
    doc.add_paragraph(f"填報日期：{datetime.datetime.now().strftime('%Y/%m/%d')}")
    
    h1 = doc.add_heading('【實驗室基本資訊】', level=2)
    if h1.runs: set_font_safe(h1.runs[0])
    
    t_base = doc.add_table(rows=3, cols=2)
    t_base.style = 'Table Grid'
    t_base.cell(0, 0).text = "系所 / 老師"; t_base.cell(0, 1).text = f"{base_data['dept']} / {base_data['mgr']}"
    t_base.cell(1, 0).text = "校區 / 門牌"; t_base.cell(1, 1).text = f"{base_data['campus']} / {base_data['room']}"
    t_base.cell(2, 0).text = "電子郵件 / 分機"; t_base.cell(2, 1).text = f"{base_data['mail']} / {base_data['ext']}"
    
    for row in t_base.rows:
        for cell in row.cells: cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        
    h2 = doc.add_heading('【年度購買紀錄申報】', level=2)
    if h2.runs: set_font_safe(h2.runs[0])
    
    if status == "無購買" or not pur_data:
        doc.add_paragraph("☑️ 盤查期間無購買氣體鋼瓶")
    else:
        for idx, p in enumerate(pur_data):
            p_table = doc.add_table(rows=3, cols=2)
            p_table.style = 'Table Grid'
            p_table.cell(0, 0).text = "鋼瓶氣體種類"; p_table.cell(0, 1).text = str(p['鋼瓶氣體種類'])
            p_table.cell(1, 0).text = "購買日期"; p_table.cell(1, 1).text = str(p['購買日期'])
            p_table.cell(2, 0).text = "購買量(公斤)"; p_table.cell(2, 1).text = f"{p['年度氣體鋼瓶購買量(公斤)']} 公斤"
            
            for row in p_table.rows:
                for cell in row.cells: cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            doc.add_paragraph("")
            
        # 移至第 2 頁顯示單據與照片
        doc.add_page_break()
        p_proof_title = doc.add_paragraph()
        p_proof_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = p_proof_title.add_run("【年度購買單據佐證】")
        run_title.bold = True
        run_title.font.size = Pt(16)
        set_font_safe(run_title)
        
        for idx, p in enumerate(pur_data):
            doc.add_paragraph(f"單據 {idx+1} ({p['鋼瓶氣體種類']})：").alignment = WD_ALIGN_PARAGRAPH.CENTER
            file_obj = p['file']
            if file_obj:
                if file_obj.type in ['image/jpeg', 'image/png']:
                    try:
                        file_obj.seek(0)
                        p_img = doc.add_paragraph()
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_img.add_run().add_picture(io.BytesIO(file_obj.read()), width=Inches(5.5))
                    except Exception: doc.add_paragraph("⚠️ [圖片單據載入異常]").alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif file_obj.type == 'application/pdf':
                    img_streams = get_pdf_preview_images(file_obj)
                    if img_streams:
                        for img_idx, img_stream in enumerate(img_streams):
                            p_img = doc.add_paragraph()
                            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p_img.add_run().add_picture(img_stream, width=Inches(5.5))
                    else:
                        doc.add_paragraph(f"📄 附檔佐證單據：{file_obj.name}").alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_paragraph("⚠️ 未建立購買單據連結").alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph("")

    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

# ================= 主程式 =================
def main():
    apply_morandi_theme()
    
    is_admin = st.session_state.get("authentication_status") and st.session_state.get("username") in ["admin"]
    url_token = st.query_params.get("token", None)
    
    # 🚀 精準修正：管理者「通知批次」選單與老師篩選機制
    if not url_token: 
        if is_admin:
            st.success("👑 管理者檢視模式：可選擇批次與老師表單進行預覽及下載。")
            records = fetch_tracker_records()
            if not records:
                st.info("資料庫尚無發送紀錄。")
                st.stop()
            
            # --- 1. 抓取所有批次並建立選單 ---
            batches = list(set([r.get('發送批次', '未分類') for r in records if r.get('專屬金鑰(Token)')]))
            batches.sort(reverse=True) # 預設顯示最新
            selected_batch = st.selectbox("📌 請選擇通知批次 (按年度管理)：", batches)
            
            # --- 2. 針對該批次篩選老師 ---
            filtered_records = [r for r in records if r.get('發送批次', '未分類') == selected_batch and r.get('專屬金鑰(Token)')]
            options = {f"[{r.get('回覆狀態', '未回報')}] {r.get('系所', '')} - {r.get('實驗室老師', '')} ({r.get('氣體鋼瓶所在位置實驗室門牌', '')})": r.get('專屬金鑰(Token)', '') for r in filtered_records}
            
            if not options:
                st.warning("該批次尚無名單。")
                st.stop()
                
            selected_teacher = st.selectbox("🔍 請選擇要檢視的實驗室表單：", list(options.keys()))
            url_token = options[selected_teacher]
            
            if not url_token: st.stop()
            st.markdown("---")
        else:
            st.error("🚫 請透過專屬通知信內的連結進入系統。")
            st.stop()

    if st.session_state.get("token") != url_token:
        st.session_state.clear(); st.session_state.token = url_token
        records = fetch_tracker_records()
        record, row_idx = next(((r, i+2) for i, r in enumerate(records) if str(r.get("專屬金鑰(Token)")).strip() == url_token), (None, None))
        
        if not record: st.error("🚫 無效的專屬連結，或連結已失效。"); st.stop()
            
        st.session_state.status = str(record.get("回覆狀態", ""))
        st.session_state.dept = str(record.get("系所", ""))
        st.session_state.manager = str(record.get("實驗室老師", ""))
        st.session_state.room = str(record.get("氣體鋼瓶所在位置實驗室門牌", ""))
        st.session_state.campus = str(record.get("校區", ""))
        st.session_state.email = str(record.get("電子郵件", ""))
        st.session_state.row_idx = row_idx
        
        df_inv = pd.DataFrame(fetch_main_inventory())
        if not df_inv.empty:
            mask = (df_inv['系所'].astype(str) == st.session_state.dept) & (df_inv['實驗室老師'].astype(str) == st.session_state.manager) & (df_inv['氣體鋼瓶所在位置實驗室門牌'].astype(str) == st.session_state.room)
            st.session_state.inventory = df_inv[mask].to_dict('records')
        else: st.session_state.inventory = []
        
        st.session_state.step = "input"

    is_readonly = (st.session_state.status == "已回報")
    
    st.markdown(f"""
    <div class="title-box">
        <h2 style="margin: 0;">國立嘉義大學 氣體鋼瓶年度盤查與回報系統</h2>
    </div>
    """, unsafe_allow_html=True)
    
    if is_readonly: st.success("✅ 該實驗室已完成本期回報，目前為唯讀檢視模式。")
    if is_readonly and not is_admin: st.stop() # 管理者即使唯讀也要能預覽
    
    saved_d = st.session_state.get("form_data", {})
    def_dept = saved_d.get("dept", st.session_state.dept)
    def_mgr = saved_d.get("mgr", st.session_state.manager)
    def_campus = saved_d.get("campus", st.session_state.campus)
    def_room = saved_d.get("room", st.session_state.room)
    def_mail = saved_d.get("mail", st.session_state.email)
    def_ext = saved_d.get("ext", st.session_state.inventory[0].get("分機", "") if st.session_state.inventory else "")
    def_buy_status = saved_d.get("status", "有購買")
    saved_inv = saved_d.get("inv", [])

    # ================= 第一階段：填寫資料 =================
    if st.session_state.step == "input" and (not is_readonly or is_admin):
        
        st.markdown('<div class="section-header">👩‍🔬 實驗室基本資訊確認與修改</div>', unsafe_allow_html=True)
        c1, c2, c3 = st.columns(3)
        with c1: e_dept = st.text_input("系所", value=def_dept)
        with c2: e_mgr = st.text_input("實驗室老師", value=def_mgr)
        with c3: e_campus = st.selectbox("校區", CAMPUS_OPTS, index=CAMPUS_OPTS.index(def_campus) if def_campus in CAMPUS_OPTS else 0)
        
        c4, c5, c6 = st.columns(3)
        with c4: e_room = st.text_input("實驗室門牌位置", value=def_room, help="若有多個位置可一併填寫，如A32-101")
        with c5: e_mail = st.text_input("電子郵件", value=def_mail)
        with c6: e_ext = st.text_input("聯絡分機", value=def_ext)
        
        st.markdown('<div class="section-header">🧪 溫室氣體盤查範疇之氣體鋼瓶使用種類現況</div>', unsafe_allow_html=True)
        st.info("💡 請確認下方登記的鋼瓶現況。若數量有異動請直接修改；若某種氣體已清空，請將數量改為 0。")
        
        collected_inv = []
        
        for idx, inv in enumerate(st.session_state.inventory):
            st.markdown(f"**🔹 氣體鋼瓶種類分項現況 {idx+1}**", unsafe_allow_html=True)
            cc1, cc2 = st.columns(2)
            
            if idx < len(saved_inv):
                g_type = saved_inv[idx]["鋼瓶氣體種類"]
                g_qty = saved_inv[idx]["鋼瓶數量"]
            else:
                g_type = str(inv.get("鋼瓶氣體種類", "請選擇")).strip()
                g_qty = int(inv.get("鋼瓶數量", 0) or 0)
                
            g_idx = GAS_TYPES.index(g_type) if g_type in GAS_TYPES else 0
            
            with cc1: e_gas = st.selectbox("氣體種類", GAS_TYPES, index=g_idx, key=f"ig_{idx}")
            with cc2: e_qty = st.number_input("鋼瓶數量 (支)", min_value=0, value=g_qty, key=f"iq_{idx}")
            st.markdown("<hr style='border: 1px dashed #BDC3C7;'>", unsafe_allow_html=True)
            if e_gas != "請選擇" and e_qty > 0: collected_inv.append({"鋼瓶氣體種類": e_gas, "鋼瓶數量": e_qty})

        st.markdown("##### ➕ 新增其他現有氣體種類")
        default_add_inv = max(0, len(saved_inv) - len(st.session_state.inventory))
        add_inv = st.number_input("欲額外新增的現有氣體種類數量", min_value=0, max_value=5, value=default_add_inv)
        
        for j in range(int(add_inv)):
            st.markdown(f"**🔹 新增現有庫存項目 {j+1}**", unsafe_allow_html=True)
            c1, c2 = st.columns(2)
            
            saved_extra_idx = len(st.session_state.inventory) + j
            if saved_extra_idx < len(saved_inv):
                n_type = saved_inv[saved_extra_idx]["鋼瓶氣體種類"]
                n_qty = saved_inv[saved_extra_idx]["鋼瓶數量"]
            else:
                n_type, n_qty = "請選擇", 1
                
            n_idx = GAS_TYPES.index(n_type) if n_type in GAS_TYPES else 0
            with c1: n_gas = st.selectbox("氣體種類", GAS_TYPES, index=n_idx, key=f"nig_{j}")
            with c2: n_qty = st.number_input("鋼瓶數量 (支)", min_value=1, value=n_qty, key=f"niq_{j}")
            st.markdown("<hr style='border: 1px dashed #BDC3C7;'>", unsafe_allow_html=True)
            if n_gas != "請選擇": collected_inv.append({"鋼瓶氣體種類": n_gas, "鋼瓶數量": n_qty})

        st.markdown('<div class="section-header">🧾 年度購買紀錄申報 (碳盤查活動數據)</div>', unsafe_allow_html=True)
        st.markdown("請依據本年度購買的單據登錄購買量並上傳佐證單據。 **若本年度無購買請直接選擇「無購買」**。")
        
        e_buy_status = st.radio("調查期間有無購買", ["有購買", "無購買"], index=["有購買", "無購買"].index(def_buy_status), horizontal=True)
        collected_purchases = []
        
        if e_buy_status == "有購買":
            saved_pur = saved_d.get("pur", [])
            def_pur_len = len(saved_pur) if saved_pur else 1
            add_pur = st.number_input("➕ 新增本年度購買單據筆數", min_value=1, max_value=10, value=max(1, def_pur_len))
            
            for k in range(int(add_pur)):
                st.markdown(f"**📦 購買單據明細建立 {k+1}**", unsafe_allow_html=True)
                pc1, pc2, pc3 = st.columns(3)
                
                saved_pur_item = saved_pur[k] if k < len(saved_pur) else {}
                saved_gas = saved_pur_item.get("鋼瓶氣體種類", "請選擇")
                saved_date = saved_pur_item.get("購買日期", datetime.date.today().strftime("%Y-%m-%d"))
                saved_amt = saved_pur_item.get("年度氣體鋼瓶購買量(公斤)", 0.0)
                saved_file = saved_pur_item.get("file", None)
                
                try: def_date = datetime.datetime.strptime(saved_date, "%Y-%m-%d").date()
                except: def_date = datetime.date.today()
                
                with pc1: p_gas = st.selectbox("購買氣體種類", GAS_TYPES, index=GAS_TYPES.index(saved_gas) if saved_gas in GAS_TYPES else 0, key=f"pg_{k}")
                with pc2: p_date = st.date_input("購買日期", value=def_date, key=f"pd_{k}")
                with pc3: p_amt = st.number_input("購買量 (公斤)", min_value=0.0, step=0.1, value=float(saved_amt), key=f"pa_{k}")
                
                st.markdown("**📎 上傳購買單據** (限 JPG, PNG, PDF)")
                p_file = st.file_uploader("", type=['png', 'jpg', 'jpeg', 'pdf'], key=f"pf_{k}", label_visibility="collapsed")
                
                final_file = p_file if p_file is not None else saved_file
                if p_file is None and saved_file is not None:
                    st.success(f"📌 已保留先前上傳單據：{saved_file.name} (若需替換請重新上傳)")
                    
                st.markdown("<hr style='border: 1px dashed #BDC3C7;'>", unsafe_allow_html=True)
                
                if p_gas != "請選擇" and p_amt > 0 and final_file:
                    collected_purchases.append({"鋼瓶氣體種類": p_gas, "購買日期": p_date.strftime("%Y-%m-%d"), "年度氣體鋼瓶購買量(公斤)": p_amt, "file": final_file})
                elif p_gas != "請選擇" and p_amt > 0 and not final_file:
                    st.warning(f"⚠️ 單據 {k+1} 尚未上傳佐證檔案，請補齊後再點擊下一步。")
        else:
            st.info("☑️ 您已勾選「盤查期間無購買氣體鋼瓶」，資料庫將記錄本年度無新增活動數據，請直接點擊下方按鈕前往預覽確認頁。")

        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("下一步：預覽確認資料", type="primary", use_container_width=True):
            if e_buy_status == "有購買" and not collected_purchases:
                st.error("❌ 您選擇了「有購買」，請務必填寫購買量並成功上傳對應單據檔案！")
            else:
                st.session_state.form_data = {
                    "dept": e_dept, "mgr": e_mgr, "campus": e_campus, "room": e_room,
                    "mail": e_mail, "ext": e_ext, "status": e_buy_status,
                    "inv": collected_inv, "pur": collected_purchases
                }
                st.session_state.step = "preview"
                st.rerun()

    # ================= 第二階段：預覽確認與送出 =================
    elif st.session_state.step == "preview" or (is_readonly and is_admin):
        st.markdown('<div class="section-header">🔍 申報內容預覽確認</div>', unsafe_allow_html=True)
        
        # 若管理者直接查看已回報，從資料庫重構 d
        if is_readonly and is_admin and "form_data" not in st.session_state:
            d = {
                "dept": st.session_state.dept, "mgr": st.session_state.manager,
                "campus": st.session_state.campus, "room": st.session_state.room,
                "mail": st.session_state.email, "ext": st.session_state.inventory[0].get("分機", "") if st.session_state.inventory else "",
                "status": "已回報 (詳細購買請見下方表單)", "inv": st.session_state.inventory, "pur": []
            }
        else:
            d = st.session_state.form_data
            
        inv_str = ', '.join([f"{i['鋼瓶氣體種類']} ({i['鋼瓶數量']}支)" for i in d['inv']]) if d['inv'] else '登記無現有庫存'
        status_color = '#B03A2E' if d.get('status')=='有購買' else '#27AE60'
        
        preview_html = f"""
<div style="background: #FDFBF7; border: 2px solid #9DB4AB; border-radius: 10px; padding: 25px; margin-bottom: 25px;">
<h3 style="color: #2C3E50; margin-top: 0; border-bottom: 2px solid #EEEEEE; padding-bottom: 10px;">📋 申報基本資料</h3>
<table style="width:100%; font-size:18px; line-height:2.2; border-collapse: collapse;">
<tr><td style="width:250px; white-space:nowrap; color:#7F8C8D;"><b>填報單位/老師</b></td><td style="font-weight:bold;">{d['dept']} / {d['mgr']} 老師</td></tr>
<tr><td style="width:250px; white-space:nowrap; color:#7F8C8D;"><b>校區/實驗門牌</b></td><td style="font-weight:bold;">{d['campus']} / {d['room']}</td></tr>
<tr><td style="width:250px; white-space:nowrap; color:#7F8C8D;"><b>聯絡分機/信箱</b></td><td style="font-weight:bold;">{d['ext']} / {d['mail']}</td></tr>
<tr><td style="width:250px; white-space:nowrap; color:#7F8C8D;"><b>最新現況總庫存</b></td><td><span style="background-color: #EBF5FB; padding: 4px 12px; border-radius: 4px; border-left: 4px solid #3498DB; font-weight:bold; color:#2980B9;">{inv_str}</span></td></tr>
</table>
<h3 style="color: #2C3E50; margin-top: 25px; border-top: 2px solid #EEEEEE; padding-top: 15px; margin-bottom: 0;">🧾 年度購買狀況：<span style="color: {status_color}; font-weight:900;">{d.get('status', '已回報')}</span></h3>
</div>
"""
        st.markdown(preview_html, unsafe_allow_html=True)
        
        if d.get('status') == "有購買":
            st.markdown("#### 📄 購買佐證憑證審視清單")
            for idx, p in enumerate(d['pur']):
                st.markdown(f"""
<div style="background-color: #F8F9F9; padding: 15px; border-radius: 6px; border-left: 5px solid #4A5D6E; margin-bottom: 15px; font-size: 26px;">
<b>單據 {idx+1}：</b> 氣體種類: <b>{p['鋼瓶氣體種類']}</b> | 購買日期: <b>{p['購買日期']}</b> | 申報購買量: <span style='color:#B03A2E; font-weight:bold;'>{p['年度氣體鋼瓶購買量(公斤)']} 公斤</span>
</div>
""", unsafe_allow_html=True)
                
                file_obj = p['file']
                if file_obj:
                    if file_obj.type in ['image/jpeg', 'image/png']:
                        st.image(file_obj, use_container_width=True, caption=f"單據 {idx+1} 影像畫面")
                    elif file_obj.type == 'application/pdf':
                        img_streams = get_pdf_preview_images(file_obj)
                        if img_streams:
                            for img_idx, img_stream in enumerate(img_streams):
                                st.image(img_stream, use_container_width=True, caption=f"單據 {idx+1} PDF 第 {img_idx+1} 頁完整預覽 (系統自動擷取)")
                        else:
                            st.markdown(f"📄 **佐證單據 PDF 檔案已暫存** (檔名: {file_obj.name})")
                st.markdown("<br>", unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)
        
        # 若為唯讀模式，僅顯示下載 Word 按鈕
        if is_readonly and is_admin:
            st.info("💡 此為已回報資料，若需產製 Word 檔請通知老師重新上傳暫存，或直接至後台下載總表。")
        else:
            col_btn1, col_btn2, col_btn3 = st.columns(3)
            with col_btn1:
                if st.button("⬅️ 回去修改資料", use_container_width=True):
                    st.session_state.step = "input"; st.rerun()
                    
            word_buf = create_report_docx(d, d['pur'], d['status'])
            
            def process_submission():
                with st.spinner("資料同步中，請稍候..."):
                    gc = get_gc(); sh_main = gc.open_by_key(SHEET_ID)
                    ws_inv = sh_main.worksheet('氣體鋼瓶資料庫'); ws_pur = sh_main.worksheet('氣體鋼瓶年度使用紀錄'); ws_trk = sh_main.worksheet('發送紀錄與金鑰')
                    
                    all_inv_records = ws_inv.get_all_records()
                    rows_to_del = [i + 2 for i, r in enumerate(all_inv_records) if str(r.get('系所', '')) == d['dept'] and str(r.get('實驗室老師', '')) == d['mgr'] and str(r.get('氣體鋼瓶所在位置實驗室門牌', '')) == st.session_state.room]
                    for r_idx in sorted(rows_to_del, reverse=True): safe_delete_row(ws_inv, r_idx); time.sleep(0.3)
                    
                    inv_to_append = [{"系所": d['dept'], "實驗室老師": d['mgr'], "校區": d['campus'], "氣體鋼瓶所在位置實驗室門牌": d['room'], "電子郵件": d['mail'], "分機": d['ext'], "鋼瓶氣體種類": item["鋼瓶氣體種類"], "鋼瓶數量": item["鋼瓶數量"], "建檔年度": f"{datetime.datetime.now().year-1911}年"} for item in d['inv']]
                    if inv_to_append: append_dicts(ws_inv, inv_to_append)
                    
                    pur_to_append = []
                    now_str = datetime.datetime.now().strftime("%Y/%m/%d %H:%M:%S")
                    for p in d['pur']:
                        ext = os.path.splitext(p["file"].name)[1]
                        f_name = f"{d['campus']}_{d['dept']}_{d['mgr']}_{p['鋼瓶氣體種類']}_{p['購買日期']}_{p['年度氣體鋼瓶購買量(公斤)']}{ext}"
                        drive_link = upload_to_drive(p["file"], f_name)
                        pur_to_append.append({"填報時間": now_str, "系所": d['dept'], "實驗室老師": d['mgr'], "校區": d['campus'], "氣體鋼瓶所在位置實驗室門牌": d['room'], "鋼瓶氣體種類": p["鋼瓶氣體種類"], "購買日期": p["購買日期"], "年度氣體鋼瓶購買量(公斤)": p["年度氣體鋼瓶購買量(公斤)"], "購買單據連結": drive_link})
                    if pur_to_append: append_dicts(ws_pur, pur_to_append)
                    
                    ws_trk.update_cell(st.session_state.row_idx, 11, "已回報")
                    ws_trk.update_cell(st.session_state.row_idx, 12, now_str)
                    st.session_state.status = "已回報"

            with col_btn2:
                if st.download_button("📥 正式申報並下載申報檔留存", data=word_buf, file_name=f"氣體鋼瓶申報存查表_{d['dept']}_{d['mgr']}.docx", use_container_width=True):
                    process_submission(); st.success("✅ 申報完成！"); time.sleep(1.5); st.rerun()
                    
            with col_btn3:
                if st.button("🚀 正式申報不下載檔案留存", key="btn_submit_no_dl", use_container_width=True):
                    process_submission(); st.success("✅ 申報完成！"); st.balloons(); time.sleep(1.5); st.rerun()

if __name__ == "__main__":
    main()