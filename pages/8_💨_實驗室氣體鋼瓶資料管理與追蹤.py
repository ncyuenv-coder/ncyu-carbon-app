import streamlit as st
import pandas as pd
import gspread
import uuid
import datetime
import smtplib
import time
import io
import re
import fitz  # PyMuPDF 用於將 PDF 轉為圖片
from PIL import Image, ImageOps  # 新增：ImageOps 用於校正手機相機旋轉
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from gspread.utils import rowcol_to_a1
import streamlit_authenticator as stauth

# Word 處理與排版所需套件
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ================= 系統與全域參數設定 =================
st.set_page_config(page_title="溫室氣體盤查範疇之氣體鋼瓶購買盤查與追蹤管理", page_icon="📢", layout="wide")

# ================= 身份驗證與初始化 =================
def clean_secrets(obj):
    if isinstance(obj, dict) or "AttrDict" in str(type(obj)): return {k: clean_secrets(v) for k, v in obj.items()}
    elif isinstance(obj, list): return [clean_secrets(i) for i in obj]
    return obj

try:
    _raw_creds = st.secrets["credentials"]
    credentials_login = clean_secrets(_raw_creds)
    cookie_cfg = st.secrets["cookie"]
    authenticator = stauth.Authenticate(credentials_login, cookie_cfg["name"], cookie_cfg["key"], cookie_cfg["expiry_days"])
except: pass

if st.session_state.get("authentication_status") is not True:
    st.warning("🔒 請先至首頁登入系統")
    st.stop()

username = st.session_state.get("username")
if username != 'admin':
    st.error("🚫 權限不足：此頁面僅供系統管理員使用。")
    st.stop()

SHEET_ID = '1Hw4rXo4ww7O9YXTwoUJeWioO5ZzM_bivRcLLpOl26DY'
BASE_FORM_URL = "https://ncyu-carbon-app-mduue5hffp7uknsskmjet9.streamlit.app/"

GAS_TYPES = ["二氧化碳", "甲烷", "乙炔", "一氧化二氮(笑氣)"]
CAMPUS_OPTS = ["蘭潭校區", "民雄校區", "新民校區", "林森校區"]

# ================= 莫蘭迪風格 CSS =================
def apply_morandi_theme():
    st.markdown("""
    <style>
        [data-testid="stAppViewContainer"] { background-color: #FDFBF7; color: #2C3E50; }
        [data-testid="stHeader"] { background-color: rgba(0,0,0,0); }
        
        /* 統計資訊卡：強制單行不換行 nowrap */
        .metric-card { background-color: #FFFFFF; padding: 0; border-radius: 12px; box-shadow: 0 4px 6px rgba(0,0,0,0.05); text-align: center; overflow: hidden; border: 1px solid #9DB4AB; display: flex; flex-direction: column; height: 100%; }
        .metric-title { background-color: #5C6B73; color: #FFFFFF; padding: 12px 15px; font-size: 18px; font-weight: bold; margin: 0; letter-spacing: 1px; }
        .metric-value { font-size: 48px; font-weight: bold; color: #000000 !important; padding: 15px; margin: 0; flex: 1; display: flex; align-items: baseline; justify-content: center; gap: 8px; flex-direction: row; white-space: nowrap;}
        
        /* 頁籤設計 */
        div[data-baseweb="tab-list"] { gap: 4px; }
        div[data-baseweb="tab-list"] button { background-color: #3B4A5A !important; border-radius: 8px 8px 0 0 !important; padding: 12px 25px !important; border: none !important; opacity: 1; }
        div[data-baseweb="tab-list"] button p { font-size: 20px !important; color: #FFFFFF !important; font-weight: 600 !important; }
        div[data-baseweb="tab-list"] button[aria-selected="true"] { background-color: #1A252D !important; border-top: 4px solid #E67E22 !important; border-bottom: none !important; }
        div[data-baseweb="tab-list"] button[aria-selected="true"] p { color: #E67E22 !important; }
        
        /* 全域一般按鈕莫蘭迪化與放大 */
        div.stButton > button { background-color: #8A9A8A !important; color: #FFFFFF !important; border: none !important; font-weight: bold !important; font-size: 18px !important; padding: 12px 24px !important; border-radius: 8px !important; }
        div.stButton > button:hover { background-color: #707F70 !important; }
        div.stDownloadButton > button { background-color: #8A9A8A !important; color: #FFFFFF !important; border: none !important; font-weight: bold !important; font-size: 24px !important; padding: 12px 24px !important; border-radius: 8px !important; }
        div.stDownloadButton > button:hover { background-color: #707F70 !important; }
        button[kind="primary"] { background-color: #5C6B73 !important; }
        button[kind="primary"]:hover { background-color: #4A565C !important; }
        
        /* 測試/正式寄信模式 Radio 按鈕莫蘭迪化 */
        div.stRadio > div[role="radiogroup"] { display: flex; gap: 15px; }
        div.stRadio > div[role="radiogroup"] > label { background-color: #E4E9E5 !important; padding: 10px 20px !important; border-radius: 8px !important; border: 2px solid #C4CDC5 !important; cursor: pointer; }
        div.stRadio > div[role="radiogroup"] > label p { color: #2C3E50 !important; font-size: 18px !important; font-weight: bold !important; margin: 0; }
        div.stRadio > div[role="radiogroup"] > label[data-checked="true"] { background-color: #9DB4AB !important; border-color: #5C6B73 !important; }
        div.stRadio > div[role="radiogroup"] > label[data-checked="true"] p { color: #FFFFFF !important; }
        
        /* 點擊展開欄 (Expander) */
        [data-testid="stExpander"] details summary { background-color: #5C6B73 !important; border-radius: 6px; padding: 12px 15px !important; }
        [data-testid="stExpander"] details summary p { font-size: 20px !important; color: #FFFFFF !important; font-weight: bold !important; margin: 0 !important; }
        [data-testid="stExpander"] details summary svg { color: #FFFFFF !important; }
    </style>
    """, unsafe_allow_html=True)

# ================= Google API 與函式 =================
def with_retry(max_retries=3, base_delay=1.5):
    def decorator(func):
        def wrapper(*args, **kwargs):
            retries = 0
            while True:
                try: return func(*args, **kwargs)
                except Exception as e:
                    if retries >= max_retries: raise e
                    time.sleep(base_delay * (2 ** retries)); retries += 1
        return wrapper
    return decorator

@st.cache_resource
def get_credentials():
    oauth_secrets = st.secrets["gcp_oauth"]
    return Credentials(token=None, refresh_token=oauth_secrets["refresh_token"], client_id=oauth_secrets["client_id"], client_secret=oauth_secrets["client_secret"], token_uri="https://oauth2.googleapis.com/token")

def get_gc(): return gspread.authorize(get_credentials())

@st.cache_data(ttl=60, show_spinner=False)
def load_data():
    gc = get_gc(); sh = gc.open_by_key(SHEET_ID)
    try: df_inv = pd.DataFrame(sh.worksheet('氣體鋼瓶資料庫').get_all_records())
    except: df_inv = pd.DataFrame()
    try: df_pur = pd.DataFrame(sh.worksheet('氣體鋼瓶年度使用紀錄').get_all_records())
    except: df_pur = pd.DataFrame()
    try: df_trk = pd.DataFrame(sh.worksheet('發送紀錄與金鑰').get_all_records())
    except: df_trk = pd.DataFrame()
    return df_inv, df_pur, df_trk

@with_retry()
def safe_append_rows(worksheet, rows): worksheet.append_rows(rows)

@with_retry()
def safe_delete_row(worksheet, idx): worksheet.delete_rows(idx)

@with_retry()
def batch_update_safe(worksheet, updates): worksheet.batch_update(updates)

def send_email_action(to_email, subject, html_body):
    try:
        sender_email = st.secrets["smtp"]["email"]; app_password = st.secrets["smtp"]["password"]; smtp_server = st.secrets["smtp"]["server"]; smtp_port = st.secrets["smtp"]["port"]
        msg = MIMEMultipart('alternative'); msg['Subject'] = subject; msg['From'] = sender_email; msg['To'] = to_email
        msg.attach(MIMEText(html_body, 'html'))
        server = smtplib.SMTP(smtp_server, smtp_port); server.starttls(); server.login(sender_email, app_password)
        server.send_message(msg); server.quit(); return True
    except Exception as e: return False

def generate_styled_email_html(email_body_text, title="溫室氣體盤查 氣體鋼瓶通知"):
    html_content = email_body_text.replace("\n", "<br>")
    return f"<html><body style='font-family: Arial, sans-serif; line-height: 1.6; color: #333; font-size: 15px;'><div style='max-width: 600px; margin: 0 auto; border: 1px solid #ddd; border-radius: 10px; overflow: hidden;'><div style='background-color: #5C6B73; color: white; padding: 15px 20px; text-align: center;'><h2 style='margin: 0;'>{title}</h2></div><div style='padding: 20px;'>{html_content}</div></div></body></html>"

@st.cache_data(ttl=600, show_spinner="產製佐證資料 Word 檔中 (包含單據圖片/PDF下載)...")
def cached_create_proof_word(df_pur_dict, current_year):
    df_pur = pd.DataFrame(df_pur_dict)
    doc = Document()
    
    # --- 樣式與邊界設定 ---
    # 設定中等邊界 (上下2.54cm, 左右1.91cm)
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(1.91)
        section.right_margin = Cm(1.91)
        
    # 設定全域字型 (英數 Times New Roman, 中文 標楷體)
    for style_name in ['Normal', 'Heading 1', 'Heading 2', 'Title']:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = 'Times New Roman'
            style._element.rPr.rFonts.set(qn('w:eastAsia'), '標楷體')

    # 表頭標題與水平置中
    h = doc.add_heading(f'{current_year}年度氣體鋼瓶購買單據佐證資料', 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    df_pur['校區'] = df_pur.get('校區', '未分類校區')
    df_pur['購買量_數值'] = pd.to_numeric(df_pur['年度氣體鋼瓶購買量(公斤)'], errors='coerce').fillna(0)
    df_pur = df_pur[df_pur['購買量_數值'] > 0].copy()
    
    if df_pur.empty:
        doc.add_paragraph("該年度無購買紀錄。")
        buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

    df_pur.sort_values(by=['校區', '系所'], inplace=True)

    doc.add_heading('【各校區氣體購買量總計】', level=1)
    summary_df = df_pur.groupby(['校區', '鋼瓶氣體種類'])['購買量_數值'].sum().reset_index()
    for campus in summary_df['校區'].unique():
        p = doc.add_paragraph()
        p.add_run(f"📍 {campus}：\n").bold = True
        c_data = summary_df[summary_df['校區'] == campus]
        for _, r in c_data.iterrows():
            p.add_run(f"    - {r['鋼瓶氣體種類']}：{r['購買量_數值']} kg\n")
    doc.add_page_break()

    try: creds = get_credentials(); drive_service = build('drive', 'v3', credentials=creds); can_dl = True
    except: can_dl = False
        
    for idx, row in df_pur.iterrows():
        doc.add_heading(f"紀錄：{row.get('校區', '')} - {row.get('系所', '')} - {row.get('實驗室老師', '')}", level=2)
        p = doc.add_paragraph()
        p.add_run(f"氣體種類：{row.get('鋼瓶氣體種類', '')}\n").bold = True
        
        date_str = row.get('購買日期', '')
        if date_str: date_str = pd.to_datetime(date_str).strftime('%Y-%m-%d')
        p.add_run(f"購買日期：{date_str}\n")
        p.add_run(f"購買量：{row.get('購買量_數值', 0)} kg\n")
        
        link = str(row.get('購買單據連結', ''))
        if link and 'drive.google.com' in link and can_dl:
            match = re.search(r'/d/([a-zA-Z0-9_-]+)', link)
            if match:
                file_id = match.group(1)
                try:
                    meta = drive_service.files().get(fileId=file_id, fields='mimeType, name').execute()
                    mime = meta.get('mimeType', '')
                    
                    if mime.startswith('image/') or mime == 'application/pdf':
                        p_img = doc.add_paragraph()
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run_img = p_img.add_run()
                        
                        if mime.startswith('image/'):
                            request = drive_service.files().get_media(fileId=file_id)
                            img_bytes = request.execute()
                            try:
                                # 1. 讀取並校正手機 EXIF 旋轉 (確保直橫判定絕對正確)
                                img = Image.open(io.BytesIO(img_bytes))
                                img = ImageOps.exif_transpose(img)
                                w, h = img.size
                                
                                # 2. 轉為乾淨的 Bytes 給 Word
                                corrected_bytes = io.BytesIO()
                                img_format = img.format if img.format else 'JPEG'
                                if img_format == 'JPEG' and img.mode in ('RGBA', 'P'):
                                    img = img.convert('RGB')
                                img.save(corrected_bytes, format=img_format)
                                corrected_bytes.seek(0)
                                
                                # 3. 計算圖片物理尺寸 (防呆：無 DPI 資訊則預設 96)
                                dpi_x, dpi_y = img.info.get('dpi', (96, 96))
                                if dpi_x == 0 or dpi_y == 0: dpi_x, dpi_y = 96, 96
                                phys_h_cm = (h / dpi_y) * 2.54
                                phys_w_cm = (w / dpi_x) * 2.54
                                
                                # 4. 嚴謹判斷直/橫式並限制高度 (同時防止寬度超過邊界 17cm)
                                if h > w: # 直式照片
                                    if phys_h_cm > 18.0:
                                        target_w = 18.0 * (w / h)
                                        if target_w > 17.0: run_img.add_picture(corrected_bytes, width=Cm(17.0))
                                        else: run_img.add_picture(corrected_bytes, height=Cm(18.0))
                                    else:
                                        if phys_w_cm > 17.0: run_img.add_picture(corrected_bytes, width=Cm(17.0))
                                        else: run_img.add_picture(corrected_bytes)
                                else: # 橫式或正方形照片
                                    if phys_h_cm > 10.0:
                                        target_w = 10.0 * (w / h)
                                        if target_w > 17.0: run_img.add_picture(corrected_bytes, width=Cm(17.0))
                                        else: run_img.add_picture(corrected_bytes, height=Cm(10.0))
                                    else:
                                        if phys_w_cm > 17.0: run_img.add_picture(corrected_bytes, width=Cm(17.0))
                                        else: run_img.add_picture(corrected_bytes)
                            except Exception as e:
                                # 若處理失敗的備案：統一塞入預設寬度
                                run_img.add_picture(io.BytesIO(img_bytes), width=Cm(15.0))
                        elif mime == 'application/pdf':
                            try:
                                request = drive_service.files().get_media(fileId=file_id)
                                pdf_stream = io.BytesIO(request.execute())
                                doc_pdf = fitz.open(stream=pdf_stream, filetype="pdf")
                                if len(doc_pdf) > 0:
                                    page = doc_pdf.load_page(0)
                                    pix = page.get_pixmap(dpi=150)
                                    img_stream = io.BytesIO(pix.tobytes("png"))
                                    run_img.add_picture(img_stream, height=Cm(18.0)) # PDF統一高度 18 公分
                                else:
                                    doc.add_paragraph("⚠️ PDF檔為空或無法讀取")
                            except Exception as e:
                                doc.add_paragraph(f"⚠️ PDF轉檔圖片失敗，請點擊連結檢視：\n{link}")
                    else: doc.add_paragraph(f"📎 附檔非圖片/PDF格式 ({meta.get('name', '未命名')})，請點擊下方連結檢視：\n{link}")
                except Exception: doc.add_paragraph(f"📎 檔案讀取限制，請點擊連結檢視：\n{link}")
        elif link: doc.add_paragraph(f"📎 單據連結：\n{link}")
        else: doc.add_paragraph("⚠️ 無購買單據連結")
        doc.add_page_break()
        
    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()

@st.fragment
def render_dashboard(df_inv, df_pur):
    st.markdown("<br>", unsafe_allow_html=True)
    all_years = []
    if not df_pur.empty:
        df_pur['購買日期'] = pd.to_datetime(df_pur['購買日期'], errors='coerce')
        df_pur['年份'] = df_pur['購買日期'].dt.year.fillna(datetime.datetime.now().year).astype(int)
        all_years = sorted(df_pur['年份'].unique().tolist(), reverse=True)
    if not all_years: all_years = [datetime.datetime.now().year]
    
    c_year, _ = st.columns([1, 3])
    sel_year = c_year.selectbox("📅 檢視年度", all_years, key="dash_year")
    st.markdown("---")
    
    df_year = df_pur[df_pur['年份'] == sel_year] if not df_pur.empty else pd.DataFrame()
    
    tot_labs = df_inv['氣體鋼瓶所在位置實驗室門牌'].nunique() if not df_inv.empty else 0
    gas_counts = df_inv.groupby('鋼瓶氣體種類')['氣體鋼瓶所在位置實驗室門牌'].nunique().to_dict() if not df_inv.empty else {}
    
    # 莫蘭迪橘色 #C06C47
    gas_summary_items = [f'<span style="font-size: 20px; font-weight: bold;">{k}</span> <span style="font-size: 48px; color: #C06C47;">{v}</span> <span style="font-size: 20px;">間</span>' for k, v in gas_counts.items()]
    gas_summary_html = ' <span style="font-size: 30px; color: #9DB4AB; margin: 0 15px;">｜</span> '.join(gas_summary_items)
    
    co2_kg = pd.to_numeric(df_year[df_year['鋼瓶氣體種類'] == '二氧化碳']['年度氣體鋼瓶購買量(公斤)'], errors='coerce').fillna(0).sum() if not df_year.empty else 0
    acet_kg = pd.to_numeric(df_year[df_year['鋼瓶氣體種類'] == '乙炔']['年度氣體鋼瓶購買量(公斤)'], errors='coerce').fillna(0).sum() if not df_year.empty else 0
    
    c1, c2, c3, c4 = st.columns(4)
    c1.markdown(f"""<div class="metric-card"><div class="metric-title">盤查範疇之氣體鋼瓶實驗室總數量</div><div class="metric-value"><span style="white-space: nowrap;"><span style="font-size: 48px; color: #C06C47;">{tot_labs}</span> <span style="font-size: 20px;">間</span></span></div></div>""", unsafe_allow_html=True)
    c2.markdown(f"""<div class="metric-card"><div class="metric-title">盤查範疇之氣體鋼瓶種類及實驗室數量</div><div class="metric-value"><span style="white-space: nowrap;">{gas_summary_html}</span></div></div>""", unsafe_allow_html=True)
    c3.markdown(f"""<div class="metric-card"><div class="metric-title">二氧化碳鋼瓶年度購買量</div><div class="metric-value"><span style="white-space: nowrap;"><span style="font-size: 48px; color: #C06C47;">{co2_kg:,.2f}</span> <span style="font-size: 20px;">公斤</span></span></div></div>""", unsafe_allow_html=True)
    c4.markdown(f"""<div class="metric-card"><div class="metric-title">乙炔鋼瓶年度購買量</div><div class="metric-value"><span style="white-space: nowrap;"><span style="font-size: 48px; color: #C06C47;">{acet_kg:,.2f}</span> <span style="font-size: 20px;">公斤</span></span></div></div>""", unsafe_allow_html=True)
    
    st.markdown("---")
    
    if not df_inv.empty:
        for dept in sorted(df_inv['系所'].astype(str).unique()):
            dept_inv = df_inv[df_inv['系所'] == dept]
            # 展開區塊數量標記 (改為純淨白色顯示)
            with st.expander(f"📁 {dept} (共 {dept_inv['氣體鋼瓶所在位置實驗室門牌'].nunique()} 間實驗室)"):
                st.markdown("<br>", unsafe_allow_html=True)
                for idx, row in dept_inv.iterrows():
                    mgr = row.get('實驗室老師', ''); room = row.get('氣體鋼瓶所在位置實驗室門牌', ''); gas = row.get('鋼瓶氣體種類', '')
                    qty = row.get('鋼瓶數量', 1); email = row.get('電子郵件', '-'); ext = row.get('分機', '-'); campus = row.get('校區', '-')
                    
                    pur_record = pd.DataFrame()
                    if not df_year.empty: pur_record = df_year[(df_year['實驗室老師'] == mgr) & (df_year['氣體鋼瓶所在位置實驗室門牌'] == room) & (df_year['鋼瓶氣體種類'] == gas)]
                    
                    has_pur = not pur_record.empty and pd.to_numeric(pur_record.iloc[0].get('年度氣體鋼瓶購買量(公斤)', 0), errors='coerce') > 0
                    pur_qty = pur_record.iloc[0].get('年度氣體鋼瓶購買量(公斤)', '0.0') if has_pur else '0.0'
                    raw_link = pur_record.iloc[0].get('購買單據連結', '') if has_pur else ''
                    report_count = 1 if has_pur else 0
                    
                    link_html = f'<a href="{raw_link}" target="_blank" style="background-color: #8A9A8A; color: white; padding: 6px 14px; border-radius: 6px; text-decoration: none; font-weight: bold; font-size: 15px;">📎 購買單據連結佐證</a>' if raw_link else '<span style="background-color: #FADBD8; color: #922B21; padding: 6px 12px; border-radius: 6px; font-size: 15px; font-weight: bold;">⚠️ 尚未申報</span>'
                    
                    st.markdown(f"""<div style="border: 1px solid #DCE4E3; border-radius: 10px; overflow: hidden; margin-bottom: 25px; box-shadow: 0 4px 8px rgba(0,0,0,0.05); background: #FDFBF7;"><div style="background-color: #D5C6E0; padding: 15px 25px; display: flex; justify-content: space-between; align-items: center; border-bottom: 2px solid #C4B5D0;"><div style="font-size: 22px; font-weight: 900; color: #000000;">{mgr} 老師</div><div style="text-align: right; display: flex; align-items: baseline; gap: 8px; white-space: nowrap;"><span style="font-size: 20px; font-weight: 900; color: #000000;">{gas}</span><span style="font-size: 28px; font-weight: 900; color: #B03A2E;">{pur_qty}</span><span style="font-size: 18px; color: #333;">公斤</span></div></div><div style="background-color: #FFFFFF;"><div style="display: flex; padding: 20px 0; border-bottom: 1px solid #EEEEEE; flex-wrap: wrap;"><div style="flex: 1 1 12%; text-align: center; border-right: 1px solid #EEEEEE;"><div style="font-size: 16px; color: #7F8C8D; margin-bottom: 6px;">校區</div><div style="font-size: 18px; font-weight: bold; color: #2C3E50;">{campus}</div></div><div style="flex: 1 1 16%; text-align: center; border-right: 1px solid #EEEEEE;"><div style="font-size: 16px; color: #7F8C8D; margin-bottom: 6px;">所屬部門</div><div style="font-size: 18px; font-weight: bold; color: #2C3E50;">{dept}</div></div><div style="flex: 1 1 12%; text-align: center; border-right: 1px solid #EEEEEE;"><div style="font-size: 16px; color: #7F8C8D; margin-bottom: 6px;">分機</div><div style="font-size: 18px; font-weight: bold; color: #2C3E50;">{ext}</div></div><div style="flex: 1 1 26%; text-align: center; border-right: 1px solid #EEEEEE;"><div style="font-size: 16px; color: #7F8C8D; margin-bottom: 6px;">電子郵件</div><div style="font-size: 18px; font-weight: bold; color: #2C3E50; word-break: break-all;">{email}</div></div><div style="flex: 1 1 18%; text-align: center; border-right: 1px solid #EEEEEE;"><div style="font-size: 16px; color: #7F8C8D; margin-bottom: 6px;">實驗室門牌</div><div style="font-size: 18px; font-weight: bold; color: #2C3E50;">{room}</div></div><div style="flex: 1 1 12%; text-align: center;"><div style="font-size: 16px; color: #7F8C8D; margin-bottom: 6px;">鋼瓶數量</div><div style="font-size: 18px; font-weight: bold; color: #2C3E50;">{qty}</div></div></div><div style="background-color: #F8F9F9; padding: 15px 25px; display: flex; justify-content: space-between; align-items: center;"><div style="font-size: 18px; color: #555;">年度申報次數：<b>{report_count}次</b></div><div>{link_html}</div></div></div></div>""", unsafe_allow_html=True)

# ================= 主程式 =================
def main():
    apply_morandi_theme()
    st.markdown('<div style="font-size: 2.4rem; font-weight: 900; color: #2C3E50; margin-bottom: 16px;">👑 溫室氣體盤查範疇之氣體鋼瓶購買盤查與追蹤管理</div>', unsafe_allow_html=True)
    
    if 'reset_key' not in st.session_state: st.session_state.reset_key = 0
    rk = st.session_state.reset_key
    df_inv, df_pur, df_trk = load_data()
    now_year_roc = datetime.datetime.now().year - 1911
    
    col_refresh = st.columns([8, 2])
    with col_refresh[1]:
        if st.button("🔄 強制刷新最新資料", use_container_width=True):
            load_data.clear()
            cached_create_proof_word.clear() # 加上這行：強制清除舊 Word 檔暫存
            st.rerun()

    tab1, tab2, tab3, tab4 = st.tabs(["📧 批次發送作業", "📊 回報進度追蹤與稽催", "📈 盤查範疇之氣體鋼瓶資料庫及年度購買量管理", "🛠️ 庫存資料管理"])

    with tab1:
        st.markdown("### 🔄 建立新一期盤查通知")
        if df_inv.empty: st.warning("⚠️ 氣體鋼瓶資料庫尚無任何紀錄，無法執行發送作業。")
        else:
            labs = df_inv[['系所', '實驗室老師', '電子郵件', '氣體鋼瓶所在位置實驗室門牌', '校區']].drop_duplicates()
            all_teachers = labs['實驗室老師'].unique().tolist()
            selected_teachers = st.multiselect("🎯 選擇特定發送對象 (若留空則全選批次發送)", all_teachers)
            if selected_teachers: labs = labs[labs['實驗室老師'].isin(selected_teachers)]; st.info(f"將針對選定的 {len(selected_teachers)} 位老師發送。")

            batch_name = st.text_input("📝 批次名稱", value=f"{now_year_roc}年度溫室氣體盤查範疇之氣體鋼瓶購買調查")
            data_year_roc = st.text_input("📅 盤查年度", value=f"{now_year_roc-1}年度")
            is_test_mode = st.radio("寄送模式", ["🧪 測試寄信模式", "🚀 正式寄信模式"], horizontal=True) == "🧪 測試寄信模式"
            test_email = st.text_input("📩 測試接收信箱") if is_test_mode else ""
            
            batch_subject = st.text_input("信件主旨", value="【重要通知】國立嘉義大學{批次名稱}")
            default_body = """{老師名稱}老師 您好：

依據環境部「事業應盤查登錄溫室氣體排放量之排放源」，自115年起，本校依規定，須於每年4月30日前完成前一年度溫室氣體排放量盤查登錄作業。
因實驗室氣體鋼瓶(二氧化碳、乙炔、甲烷、笑氣)屬法定盤查範疇，請老師協助於O年O月O日前檢視與更新貴實驗室之基本資料、氣體鋼瓶庫存及{年度}購買量，並上傳{年度}之購買佐證單據，謝謝。

📍 請由實驗室專屬確認連結進入填報與登錄：
{links}
 
🔔補充說明：
．氣體鋼瓶若放置於多個實驗室，建請一併填寫，如A32-101、A32-103、A32-105。
．若有任何疑問或需要協助，歡迎與環安中心(#7137)聯繫，我們將竭誠為您服務。

敬祝
教安 順心
環安中心環保組 敬上
(本信件由系統自動發送，請勿直接回覆)"""
            batch_body = st.text_area("信件內容", value=default_body, height=350)
            
            if st.button("🚀 產生金鑰並發送通知信", type="primary"):
                if is_test_mode and not test_email: st.error("請輸入測試信箱！")
                else:
                    my_bar = st.progress(0, text="處理中...")
                    grouped = labs.groupby(['系所', '實驗室老師', '電子郵件'])
                    all_appends, success_count, total_groups = [], 0, len(grouped)
                    now_time = datetime.datetime.now().strftime("%Y/%m/%d %H:%M:%S")
                    
                    for idx, ((dept, mgr, email), group) in enumerate(grouped):
                        target_email = test_email if is_test_mode else email
                        links_html = ""
                        for _, row in group.iterrows():
                            token = uuid.uuid4().hex
                            link = f"{BASE_FORM_URL}?token={token}"
                            links_html += f'<div style="margin-bottom: 8px; background-color: #FDFBF7; padding: 10px; border-left: 5px solid #D4A373;"><b>{row["校區"]} - 門牌 {row["氣體鋼瓶所在位置實驗室門牌"]}</b>：<a href="{link}" style="color: #3498DB; font-weight: bold;">點此進入系統</a></div>'
                            all_appends.append([batch_name, dept, mgr, row["校區"], row["氣體鋼瓶所在位置實驗室門牌"], email, token, link, "", now_time, "待回覆", ""])
                        
                        clean_mgr = re.sub(r'^[a-zA-Z0-9]+\s+', '', str(mgr)).replace("老師", "").strip()
                        mail_content = batch_body.replace("{系所}", dept).replace("{老師名稱}", clean_mgr).replace("{年度}", data_year_roc).replace("{links}", links_html)
                        html_wrap = generate_styled_email_html(mail_content, title=batch_name)
                        subject_content = batch_subject.replace("{批次名稱}", batch_name)
                        
                        if send_email_action(target_email, subject_content, html_wrap):
                            success_count += 1
                            for record in all_appends[-len(group):]: record[8] = "發送成功"
                        else:
                            for record in all_appends[-len(group):]: record[8] = "發送失敗"
                        my_bar.progress((idx+1)/total_groups, text=f"發送中 ({idx+1}/{total_groups})")
                    
                    safe_append_rows(get_gc().open_by_key(SHEET_ID).worksheet('發送紀錄與金鑰'), all_appends)
                    st.success(f"✅ 完成！成功寄出 {success_count} 位老師。")
                    time.sleep(2); load_data.clear(); st.rerun()

    with tab2:
        if df_trk.empty: st.info("尚無發送追蹤紀錄。")
        else:
            batches = df_trk['發送批次'].unique().tolist()
            sel_batch = st.selectbox("📌 篩選發送批次", batches, index=len(batches)-1)
            df_cur = df_trk[df_trk['發送批次'] == sel_batch]
            pending = df_cur[df_cur['回覆狀態'] == '待回覆']
            replied = df_cur[df_cur['回覆狀態'] == '已回報']
            
            st.markdown(f"""
            <div style="display: flex; gap: 20px; margin-bottom: 25px;">
                <div class="metric-card" style="flex: 1;"><div class="metric-title">應回報總數</div><div class="metric-value"><span style="white-space: nowrap;"><span style="font-size: 48px;">{len(df_cur)}</span><span style="font-size: 20px;">間</span></span></div></div>
                <div class="metric-card" style="flex: 1; border-color: #F39C12;"><div class="metric-title" style="background-color: #F39C12;">尚未回報</div><div class="metric-value"><span style="white-space: nowrap;"><span style="font-size: 48px;">{len(pending)}</span><span style="font-size: 20px;">間</span></span></div></div>
                <div class="metric-card" style="flex: 1; border-color: #27AE60;"><div class="metric-title" style="background-color: #27AE60;">已回報完成</div><div class="metric-value"><span style="white-space: nowrap;"><span style="font-size: 48px;">{len(replied)}</span><span style="font-size: 20px;">間</span></span></div></div>
            </div>
            """, unsafe_allow_html=True)
            
            st.markdown("### ✉️ 稽催通知作業")
            follow_mode = st.radio("稽催發送模式", ["🧪 測試寄信模式", "🚀 正式寄信模式"], horizontal=True) == "🧪 測試寄信模式"
            f_test_email = st.text_input("📩 稽催測試接收信箱") if follow_mode else ""
            
            st.markdown("#### 📝 編輯稽催信件")
            f_data_year = st.text_input("盤查年度 (供信件中 {年度} 變數替換使用)", value=f"{now_year_roc-1}年度")
            f_default_subject = "【稽催提醒】國立嘉義大學{年度}碳盤查範疇氣體鋼瓶資料尚未回報，敬請於O月O日前限期填報"
            f_default_body = """{系所} {老師名稱}老師 您好：
 
環安中心前已發送「溫室氣體盤查範疇之氣體鋼瓶購買調查」通知，經查尚未收到貴實驗室之回報資料。
 
因全校溫室氣體盤查登錄作業時程緊迫，為避免延誤向環境部依法申報進度而衍生罰鍰問題，請老師協助於O年O月O日前，撥冗檢視及填報貴實驗室之氣體鋼瓶（二氧化碳、乙炔、甲烷、笑氣）庫存、{年度}購買量，並上傳{年度}之購買佐證單據。

📍 請由實驗室專屬確認連結進入填報與登錄：
{專屬填報連結區塊}
 
🔔 補充說明：
．若貴實驗室未購買或無使用上述四種列管氣體，亦請點擊連結進入系統勾選「無使用」，以利系統註記。
．若您已於近日完成填報，請直接忽略此提醒信件。
．若有任何疑問或需要協助，歡迎與環安中心 (#7137) 聯繫。
 
敬祝
教安 順心
環安中心環保組 敬上
(本信件由系統自動發送，請勿直接回覆)"""
            f_subject = st.text_input("稽催信件主旨", value=f_default_subject)
            f_body = st.text_area("稽催信件內容", value=f_default_body, height=350)
            
            f_all_teachers = pending['實驗室老師'].unique().tolist()
            f_selected_teachers = st.multiselect("🎯 選擇特定稽催對象 (若留空則依系所批次/或全域稽催)", f_all_teachers)
            if f_selected_teachers: pending = pending[pending['實驗室老師'].isin(f_selected_teachers)]
            
            st.markdown("---")
            if not pending.empty:
                # 一鍵發送全部按鈕
                if st.button("🚀 一鍵稽催全部尚未回報老師", type="primary", use_container_width=True):
                    if follow_mode and not f_test_email: st.error("請輸入測試信箱！")
                    else:
                        my_bar_f = st.progress(0, text="全面稽催處理中...")
                        grouped_all = pending.groupby(['系所', '實驗室老師', '電子郵件'])
                        f_count = 0; total_f = len(grouped_all)
                        
                        for idx, ((dept, mgr, email), group) in enumerate(grouped_all):
                            target = f_test_email if follow_mode else email
                            display_mgr = re.sub(r'^[a-zA-Z0-9]+\s+', '', str(mgr))
                            display_mgr = display_mgr.replace("老師", "").strip()
                            links_html = "".join([f'<div style="margin-bottom: 8px; background-color: #FDFBF7; padding: 10px; border-left: 5px solid #F39C12;">門牌 <b>{r["氣體鋼瓶所在位置實驗室門牌"]}</b>：<a href="{r["專屬填報網址"]}" style="color: #3498DB; font-weight: bold;">點此補登</a></div>' for _, r in group.iterrows()])
                            
                            mail_content = f_body.replace("{系所}", dept).replace("{老師名稱}", display_mgr).replace("{年度}", f_data_year).replace("{專屬填報連結區塊}", links_html)
                            html_wrap = generate_styled_email_html(mail_content, title="氣體鋼瓶盤查 稽催提醒")
                            subject_content = f_subject.replace("{年度}", f_data_year)
                            
                            if send_email_action(target, subject_content, html_wrap): f_count += 1
                            my_bar_f.progress((idx+1)/total_f, text=f"發送中 ({idx+1}/{total_f})")
                        st.success(f"✅ 全面稽催發送完成 (共寄出 {f_count} 封信)。")

                # 個別系所區塊
                for dept in pending['系所'].unique():
                    dept_pending = pending[pending['系所'] == dept]
                    with st.expander(f"🏢 {dept} (未回報 {len(dept_pending)} 間)"):
                        col_list, col_btn = st.columns([8, 2])
                        grouped_pending = dept_pending.groupby(['實驗室老師', '電子郵件'])
                        with col_list:
                            for (mgr, email), group in grouped_pending:
                                ext_info = df_inv[(df_inv['系所']==dept) & (df_inv['實驗室老師']==mgr)]['分機'].iloc[0] if not df_inv.empty and len(df_inv[(df_inv['系所']==dept) & (df_inv['實驗室老師']==mgr)]) > 0 else "無資料"
                                st.write(f"- **{mgr}** 老師 (分機: {ext_info})")
                        with col_btn:
                            if st.button(f"✉️ 獨立稽催 {dept}", key=f"f_{dept}"):
                                if follow_mode and not f_test_email: st.error("請輸入測試信箱！")
                                else:
                                    f_count = 0
                                    for (mgr, email), group in grouped_pending:
                                        target = f_test_email if follow_mode else email
                                        display_mgr = re.sub(r'^[a-zA-Z0-9]+\s+', '', str(mgr))
                                        display_mgr = display_mgr.replace("老師", "").strip()
                                        links_html = "".join([f'<div style="margin-bottom: 8px; background-color: #FDFBF7; padding: 10px; border-left: 5px solid #F39C12;">門牌 <b>{r["氣體鋼瓶所在位置實驗室門牌"]}</b>：<a href="{r["專屬填報網址"]}" style="color: #3498DB; font-weight: bold;">點此補登</a></div>' for _, r in group.iterrows()])
                                        
                                        mail_content = f_body.replace("{系所}", dept).replace("{老師名稱}", display_mgr).replace("{年度}", f_data_year).replace("{專屬填報連結區塊}", links_html)
                                        html_wrap = generate_styled_email_html(mail_content, title="氣體鋼瓶盤查 稽催提醒")
                                        subject_content = f_subject.replace("{年度}", f_data_year)
                                        
                                        if send_email_action(target, subject_content, html_wrap): f_count += 1
                                    st.success(f"✅ {dept} 稽催發送完成 ({f_count} 封信)。")

    with tab3:
        render_dashboard(df_inv, df_pur)
        st.markdown("<br><hr>", unsafe_allow_html=True)
        st.markdown("### 📁 溫室氣體盤查範疇之氣體鋼瓶購買統計")
        
        if not df_pur.empty and not df_inv.empty:
            col_dl1, col_dl2 = st.columns(2)
            
            df_pur_out = df_pur.copy()
            if not df_pur_out.empty:
                df_pur_out['購買日期'] = pd.to_datetime(df_pur_out['購買日期'], errors='coerce').dt.strftime('%Y-%m-%d')
            
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
                df_pur_out.to_excel(writer, sheet_name='氣體鋼瓶年度使用紀錄', index=False)
                
            with col_dl1: st.download_button("📥 下載年度使用紀錄 Excel 檔", data=output.getvalue(), file_name=f"氣體鋼瓶盤查清冊_{datetime.datetime.now().year}.xlsx", type="primary", use_container_width=True)
            with col_dl2:
                word_bytes = cached_create_proof_word(df_pur.to_dict('records'), datetime.datetime.now().year)
                st.download_button("📄 產製並下載年度購買單據佐證資料 (Word)", data=word_bytes, file_name=f"購買單據佐證資料_{datetime.datetime.now().year}.docx", type="primary", use_container_width=True)

    with tab4:
        st.markdown("### 🛠️ 氣體鋼瓶資料庫管理")
        manage_mode = st.radio("選擇管理模式", ["➕ 新增實驗室與庫存", "🔄 現有庫存查詢與異動"], horizontal=True)
        gc = get_gc(); ws_inv = gc.open_by_key(SHEET_ID).worksheet('氣體鋼瓶資料庫')
        
        if manage_mode == "➕ 新增實驗室與庫存":
            with st.container():
                st.markdown('<div style="background-color: #8A9A8A; padding: 12px 20px; border-radius: 8px; margin-bottom: 20px;"><h3 style="color: #FFFFFF; margin: 0; font-weight: 600;">🏢 場所基本資料建立</h3></div>', unsafe_allow_html=True)
                col1, col2, col3 = st.columns(3)
                with col1: a_dept = st.text_input("系所", key=f"a_dept_{rk}")
                with col2: a_campus = st.selectbox("校區", CAMPUS_OPTS, key=f"a_camp_{rk}")
                with col3: a_room = st.text_input("氣體鋼瓶所在位置實驗室門牌", key=f"a_room_{rk}")
                col4, col5, col6 = st.columns(3)
                with col4: a_mgr = st.text_input("實驗室老師", key=f"a_mgr_{rk}")
                with col5: a_mail = st.text_input("電子郵件", key=f"a_mail_{rk}")
                with col6: a_ext = st.text_input("分機", key=f"a_ext_{rk}")
            
            with st.container():
                st.markdown('<div style="background-color: #A393B3; padding: 12px 20px; border-radius: 8px; margin-bottom: 20px; margin-top: 20px;"><h3 style="color: #FFFFFF; margin: 0; font-weight: 600;">💨 氣體鋼瓶初始庫存</h3></div>', unsafe_allow_html=True)
                num_gases = st.number_input("🧪 此場所氣體種類數量", min_value=1, max_value=5, value=1, step=1, key=f"a_num_{rk}")
                gas_entries = []
                for i in range(int(num_gases)):
                    st.markdown(f"#### 🔹 第 {i+1} 項氣體")
                    gc1, gc2, gc3 = st.columns(3)
                    with gc1: a_gas = st.selectbox("鋼瓶氣體種類", ["請選擇"] + GAS_TYPES, key=f"a_gas_{i}_{rk}")
                    with gc2: a_qty = st.number_input("鋼瓶數量", min_value=1, step=1, key=f"a_qty_{i}_{rk}")
                    with gc3: a_year = st.text_input("建檔年度", value=f"{datetime.datetime.now().year}年", key=f"a_yr_{i}_{rk}")
                    gas_entries.append({"gas": a_gas, "qty": a_qty, "year": a_year})
                    st.markdown("<hr style='border-top: 1px dashed #B8A9C9; margin: 15px 0;'>", unsafe_allow_html=True)
            
            if st.button("🚀 確認新增建檔", type="primary", use_container_width=True):
                if not a_dept or not a_mgr or not a_room: st.error("請確實填寫系所、老師與門牌！")
                else:
                    rows_to_append = []
                    for entry in gas_entries:
                        if entry["gas"] != "請選擇": rows_to_append.append([a_dept, a_mgr, a_campus, a_room, a_mail, a_ext, entry["gas"], entry["qty"], entry["year"]])
                    if rows_to_append:
                        safe_append_rows(ws_inv, rows_to_append)
                        st.success(f"✅ {a_mgr} 老師的庫存已建檔！"); st.session_state.reset_key += 1; time.sleep(1); load_data.clear(); st.rerun()
                    else: st.error("請至少選擇一種氣體！")

        elif manage_mode == "🔄 現有庫存查詢與異動":
            if df_inv.empty: st.info("目前資料庫為空。")
            else:
                existing_depts = df_inv['系所'].astype(str).unique().tolist()
                f_col1, f_col2 = st.columns(2)
                with f_col1: sel_dept = st.selectbox("1. 篩選系所", existing_depts)
                dept_labs = df_inv[df_inv['系所'] == sel_dept]
                lab_opts = dept_labs.apply(lambda x: f"{x['實驗室老師']}({x['氣體鋼瓶所在位置實驗室門牌']})", axis=1).unique().tolist()
                with f_col2: sel_lab_opt = st.selectbox("2. 篩選老師與門牌", lab_opts)
                
                mgr_sel = sel_lab_opt.split("(")[0]; room_sel = sel_lab_opt.split("(")[1].replace(")", "")
                target_df = df_inv[(df_inv['系所'] == sel_dept) & (df_inv['實驗室老師'] == mgr_sel) & (df_inv['氣體鋼瓶所在位置實驗室門牌'] == room_sel)]
                
                if not target_df.empty:
                    base_row = target_df.iloc[0]
                    campus_val = base_row['校區']
                    st.markdown('<div style="background-color: #8A9A8A; padding: 12px 20px; border-radius: 8px; margin-bottom: 20px; margin-top: 20px;"><h3 style="color: #FFFFFF; margin: 0; font-weight: 600;">🏢 場所基本資料 (編輯)</h3></div>', unsafe_allow_html=True)
                    ce1, ce2, ce3 = st.columns(3)
                    with ce1: u_dept = st.text_input("系所", value=sel_dept, key="u_dept")
                    with ce2: u_camp = st.selectbox("校區", CAMPUS_OPTS, index=CAMPUS_OPTS.index(campus_val) if campus_val in CAMPUS_OPTS else 0, key="u_camp")
                    with ce3: u_room = st.text_input("氣體鋼瓶所在位置實驗室門牌", value=room_sel, key="u_room")
                    ce4, ce5, ce6 = st.columns(3)
                    with ce4: u_mgr = st.text_input("實驗室老師", value=mgr_sel, key="u_mgr")
                    with ce5: e_mail = st.text_input("電子郵件", value=str(base_row['電子郵件']), key=f"e_m_{mgr_sel}_{room_sel}")
                    with ce6: e_ext = st.text_input("分機", value=str(base_row['分機']), key=f"e_e_{mgr_sel}_{room_sel}")
                    
                    st.markdown('<div style="background-color: #A393B3; padding: 12px 20px; border-radius: 8px; margin-bottom: 20px; margin-top: 20px;"><h3 style="color: #FFFFFF; margin: 0; font-weight: 600;">💨 氣體鋼瓶現有庫存與異動</h3></div>', unsafe_allow_html=True)
                    collected_updates = []; all_rows = ws_inv.get_all_records()
                    
                    for idx, row in target_df.iterrows():
                        gas_name = row['鋼瓶氣體種類']
                        orig_idx = next((i + 2 for i, r in enumerate(all_rows) if str(r['系所'])==sel_dept and str(r['實驗室老師'])==mgr_sel and str(r['氣體鋼瓶所在位置實驗室門牌'])==room_sel and str(r['鋼瓶氣體種類'])==gas_name), -1)
                        uc1, uc2, uc3 = st.columns([2, 2, 1])
                        with uc1: st.text_input(f"氣體種類 (現有)", value=gas_name, disabled=True, key=f"ugn_{idx}")
                        with uc2: u_qty = st.number_input("鋼瓶數量", value=int(row['鋼瓶數量']), min_value=0, key=f"uq_{idx}_{mgr_sel}_{room_sel}")
                        with uc3: st.markdown("<br>", unsafe_allow_html=True); del_flag = st.checkbox("🗑️ 刪除", key=f"del_{idx}_{mgr_sel}_{room_sel}")
                        collected_updates.append({"orig_idx": orig_idx, "gas": gas_name, "qty": u_qty, "delete": del_flag})
                        
                    st.markdown("<hr style='border-top: 1px dashed #B8A9C9; margin: 15px 0;'>", unsafe_allow_html=True)
                    st.markdown("#### ➕ 新增其他氣體種類至此實驗室")
                    ac1, ac2, ac3 = st.columns(3)
                    existing_gases = target_df['鋼瓶氣體種類'].tolist()
                    available_gases = [g for g in GAS_TYPES if g not in existing_gases]
                    
                    with ac1: new_gas = st.selectbox("選擇新增氣體", ["請選擇"] + available_gases, key=f"new_gas_{mgr_sel}_{room_sel}")
                    with ac2: new_qty = st.number_input("新增氣體數量", min_value=1, step=1, key=f"new_qty_{mgr_sel}_{room_sel}")
                    with ac3: new_year = st.text_input("建檔年度", value=f"{datetime.datetime.now().year}年", key=f"new_yr_{mgr_sel}_{room_sel}")
                        
                    st.markdown("<br>", unsafe_allow_html=True)
                    if st.button("💾 儲存所有異動", type="primary", use_container_width=True):
                        with st.spinner("同步至雲端資料庫中..."):
                            batch_data = []; rows_to_delete = []; rows_to_append = []
                            for u in collected_updates:
                                if u["orig_idx"] != -1:
                                    if u["delete"]: rows_to_delete.append(u["orig_idx"])
                                    else:
                                        batch_data.append({'range': f"A{u['orig_idx']}", 'values': [[str(u_dept)]]})
                                        batch_data.append({'range': f"B{u['orig_idx']}", 'values': [[str(u_mgr)]]})
                                        batch_data.append({'range': f"C{u['orig_idx']}", 'values': [[str(u_camp)]]})
                                        batch_data.append({'range': f"D{u['orig_idx']}", 'values': [[str(u_room)]]})
                                        batch_data.append({'range': f"E{u['orig_idx']}", 'values': [[str(e_mail)]]})
                                        batch_data.append({'range': f"F{u['orig_idx']}", 'values': [[str(e_ext)]]})
                                        batch_data.append({'range': f"H{u['orig_idx']}", 'values': [[u["qty"]]]})
                            
                            if new_gas != "請選擇":
                                rows_to_append.append([str(u_dept), str(u_mgr), str(u_camp), str(u_room), str(e_mail), str(e_ext), new_gas, new_qty, new_year])

                            if batch_data: batch_update_safe(ws_inv, batch_data)
                            for r_idx in sorted(rows_to_delete, reverse=True): safe_delete_row(ws_inv, r_idx)
                            if rows_to_append: safe_append_rows(ws_inv, rows_to_append)
                                
                            st.success("✅ 實驗室資料與庫存已成功更新！")
                            st.session_state.reset_key += 1; time.sleep(1); load_data.clear(); st.rerun()

if __name__ == "__main__":
    main()