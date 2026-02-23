import streamlit as st
import pandas as pd
from datetime import datetime, date, timedelta
import gspread
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
import streamlit_authenticator as stauth
import plotly.express as px
import plotly.graph_objects as go
import time
import re
import io

try:
    from docx import Document
    from docx.shared import Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_ORIENT
    import fitz  # PyMuPDF
    DOCX_READY = True
except ImportError:
    DOCX_READY = False

# ==========================================
# 0. 系統設定 (後台管理專用)
# ==========================================
st.set_page_config(page_title="燃油設備後台管理", page_icon="⚙️", layout="wide")

# ==========================================
# 1. CSS 樣式表
# ==========================================
st.markdown("""
<style>
    :root {
        color-scheme: light;
        --orange-bg: #E67E22;     
        --orange-dark: #D35400;
        --text-main: #2C3E50;
        --text-sub: #566573;
        --morandi-red: #C0392B; 
        --kpi-gas: #52BE80;
        --kpi-diesel: #F4D03F;
        --kpi-total: #5DADE2;
        --kpi-co2: #AF7AC5;
        --morandi-blue: #34495E;
        --deep-gray: #333333;
    }
    [data-testid="stAppViewContainer"] { background-color: #EAEDED; color: var(--text-main); }
    [data-testid="stHeader"] { background-color: rgba(0,0,0,0); }
    [data-testid="stSidebar"] { background-color: #FFFFFF; border-right: 1px solid #BDC3C7; }
    div[data-baseweb="input"] > div, div[data-baseweb="base-input"] > input, textarea, input { background-color: #FFFFFF !important; border-color: #BDC3C7 !important; color: #000000 !important; font-size: 1.15rem !important; }
    div[data-baseweb="select"] > div { border-color: #BDC3C7 !important; background-color: #FFFFFF !important; }
    ul[data-baseweb="menu"] { background-color: #FFFFFF !important; }
    div.stButton > button, button[kind="primary"] { background-color: var(--orange-bg) !important; color: #FFFFFF !important; border: 2px solid var(--orange-dark) !important; border-radius: 12px !important; font-size: 1.3rem !important; font-weight: 800 !important; padding: 0.7rem 1.5rem !important; box-shadow: 0 4px 6px rgba(230, 126, 34, 0.3) !important; width: 100%; }
    div.stButton > button p { color: #FFFFFF !important; } 
    div.stButton > button:hover { background-color: var(--orange-dark) !important; transform: translateY(-2px) !important; color: #FFFFFF !important; }
    button[data-baseweb="tab"] div p { font-size: 1.3rem !important; font-weight: 900 !important; color: var(--text-sub); }
    button[data-baseweb="tab"][aria-selected="true"] div p { color: #E67E22 !important; border-bottom: 3px solid #E67E22; }
    .stRadio div[role="radiogroup"] label { background-color: #D6EAF8 !important; border: 1px solid #AED6F1 !important; border-radius: 8px !important; padding: 8px 15px !important; margin-right: 10px !important; }
    .stRadio div[role="radiogroup"] label p { font-size: 1.25rem !important; font-weight: 800 !important; color: #000000 !important; }
    [data-testid="stDataFrame"] { font-size: 1.25rem !important; }
    [data-testid="stDataFrame"] div { font-size: 1.25rem !important; }
    .dev-card-v148 { background-color: #FFFFFF; border: 1px solid #BDC3C7; border-radius: 12px; overflow: hidden; box-shadow: 0 3px 6px rgba(0,0,0,0.08); margin-bottom: 20px; display: flex; flex-direction: column; }
    .dev-header { padding: 12px 15px; display: flex; justify-content: space-between; align-items: center; border-bottom: 1px solid rgba(0,0,0,0.1); }
    .dev-header-left { display: flex; flex-direction: column; gap: 3px; }
    .dev-id { font-size: 1.15rem; font-weight: 800; color: #000000 !important; opacity: 0.8; } 
    .dev-name-row { display: flex; align-items: center; gap: 8px; flex-wrap: wrap; }
    .dev-name { font-size: 1.25rem; font-weight: 900; color: #000000 !important; }
    .qty-badge { font-size: 0.85rem; background-color: rgba(255,255,255,0.9); padding: 2px 8px; border-radius: 12px; color: #2C3E50; font-weight: bold; border: 1px solid rgba(0,0,0,0.1); }
    .dev-header-right { text-align: right; display: flex; flex-direction: column; align-items: flex-end; justify-content: center; }
    .dev-vol { font-size: 1.8rem; color: #C0392B !important; font-weight: 900; line-height: 1.1; }
    .dev-unit { font-size: 0.95rem; color: var(--deep-gray); font-weight: bold; margin-left: 2px; }
    .dev-body { padding: 15px; font-size: 0.95rem; color: var(--deep-gray); display: grid; grid-template-columns: 1fr 1fr; gap: 8px; }
    .dev-item { margin-bottom: 2px; display: flex; align-items: baseline; }
    .dev-label { font-weight: 700; color: var(--deep-gray) !important; font-size: 0.95rem; margin-right: 5px; min-width: 80px; }
    .dev-val { color: var(--deep-gray) !important; font-weight: 600; font-size: 1rem; }
    .dev-footer { padding: 10px 15px; background-color: #F8F9F9; border-top: 1px solid #E5E7E9; display: flex; justify-content: space-between; align-items: center; }
    .dev-count { font-weight: 700; color: #34495E; font-size: 0.95rem; }
    .alert-status { color: #C0392B; font-weight: 900; display: flex; align-items: center; gap: 5px; background-color: #FADBD8; padding: 4px 12px; border-radius: 12px; font-size: 0.9rem; }
    .stat-card-v119 { background-color: #FFFFFF; border: 1px solid #BDC3C7; border-radius: 12px; overflow: hidden; box-shadow: 0 3px 6px rgba(0,0,0,0.08); margin-bottom: 15px; height: 100%; }
    .stat-header { padding: 15px 25px; display: flex; justify-content: space-between; align-items: center; border-bottom: 1px solid rgba(0,0,0,0.1); }
    .stat-title { font-size: 1.4rem; font-weight: 900; color: #2C3E50; } 
    .stat-count { font-size: 2.2rem; font-weight: 900; color: var(--morandi-red); }
    .stat-body-split { padding: 15px 20px; display: flex; }
    .stat-col-left { width: 50%; padding-right: 15px; border-right: 1px dashed #BDC3C7; }
    .stat-col-right { width: 50%; padding-left: 15px; }
    .stat-item { font-size: 1.1rem; color: #566573; margin-bottom: 8px; display: flex; justify-content: space-between; } 
    .stat-item-label { font-weight: bold; color: #2C3E50; }
    .stat-item-val { color: #2C3E50; font-weight: 900; }
    .top-kpi-card { background-color: #FFFFFF; border-radius: 12px; padding: 25px; text-align: center; box-shadow: 0 4px 10px rgba(0,0,0,0.1); border: 1px solid #BDC3C7; margin-bottom: 10px; }
    .top-kpi-title { font-size: 1.15rem; color: #7F8C8D; font-weight: bold; margin-bottom: 5px; }
    .top-kpi-value { font-size: 3.5rem; color: #2C3E50; font-weight: 900; line-height: 1.1; }
    .admin-kpi-card { background-color: #FFFFFF; border: 1px solid #BDC3C7; border-radius: 12px; overflow: hidden; box-shadow: 0 4px 10px rgba(0,0,0,0.1); height: 100%; text-align: center; margin-bottom: 20px; }
    .admin-kpi-header { padding: 10px; font-size: 1.2rem; font-weight: bold; color: #2C3E50; border-bottom: 1px solid rgba(0,0,0,0.1); }
    .admin-kpi-body { padding: 20px; }
    .admin-kpi-value { font-size: 2.8rem; font-weight: 900; color: #2C3E50; margin-bottom: 5px; }
    .admin-kpi-unit { font-size: 1rem; color: #7F8C8D; font-weight: normal; margin-left: 5px; }
    .admin-kpi-sub { font-size: 0.9rem; display: inline-block; padding: 2px 10px; border-radius: 15px; background-color: #F9E79F; color: #7D6608; margin-top: 5px; font-weight: bold; }
    .unreported-block { padding: 15px 20px; border-radius: 12px; margin-bottom: 20px; color: #2C3E50; box-shadow: 0 2px 6px rgba(0,0,0,0.08); border: 1px solid rgba(0,0,0,0.05); }
    .unreported-title { font-size: 1.6rem; font-weight: 900; margin-bottom: 12px; border-bottom: 2px solid rgba(0,0,0,0.1); padding-bottom: 8px; }
    .dashboard-main-title { font-size: 1.8rem; font-weight: 900; text-align: center; color: #2C3E50; margin-bottom: 20px; background-color: #F8F9F9; padding: 10px; border-radius: 10px; border: 1px solid #BDC3C7; }
    .bar-chart-box { border: 1px solid #BDC3C7; border-radius: 12px; padding: 15px; background-color: #FFFFFF; box-shadow: 0 2px 4px rgba(0,0,0,0.05); }
</style>
""", unsafe_allow_html=True)

# ==========================================
# 2. 身份驗證與初始化
# ==========================================
def clean_secrets(obj):
    if isinstance(obj, dict) or "AttrDict" in str(type(obj)): return {k: clean_secrets(v) for k, v in obj.items()}
    elif isinstance(obj, list): return [clean_secrets(i) for i in obj]
    return obj

if st.session_state.get("authentication_status") is not True:
    st.warning("🔒 請先至首頁 (Hello) 登入系統")
    st.stop()

username = st.session_state.get("username")
name = st.session_state.get("name")

# 安全管控：限制只有 admin 能夠進入
if username != 'admin':
    st.error("🚫 權限不足：此頁面僅供系統管理員使用。")
    st.stop()

try:
    _raw_creds = st.secrets["credentials"]
    credentials_login = clean_secrets(_raw_creds)
    cookie_cfg = st.secrets["cookie"]
    authenticator = stauth.Authenticate(credentials_login, cookie_cfg["name"], cookie_cfg["key"], cookie_cfg["expiry_days"])
except: pass

with st.sidebar:
    st.header(f"👤 {name} (管理員)")
    st.success("☁️ 後台連線正常")
    authenticator.logout('登出系統', 'sidebar')

# ==========================================
# 3. 資料庫連線與常數
# ==========================================
SHEET_ID = "1gqDU21YJeBoBOd8rMYzwwZ45offXWPGEODKTF6B8k-Y" 
DEVICE_ORDER = ["公務車輛(GV-1-)", "乘坐式割草機(GV-2-)", "乘坐式農用機具(GV-3-)", "鍋爐(GS-1-)", "發電機(GS-2-)", "肩背或手持式割草機、吹葉機(GS-3-)", "肩背或手持式農用機具(GS-4-)"]
DEVICE_CODE_MAP = {"GV-1": "公務車輛(GV-1-)", "GV-2": "乘坐式割草機(GV-2-)", "GV-3": "乘坐式農用機具(GV-3-)", "GS-1": "鍋爐(GS-1-)", "GS-2": "發電機(GS-2-)", "GS-3": "肩背或手持式割草機、吹葉機(GS-3-)", "GS-4": "肩背或手持式農用機具(GS-4-)"}
MORANDI_COLORS = { "公務車輛(GV-1-)": "#B0C4DE", "乘坐式割草機(GV-2-)": "#F5CBA7", "乘坐式農用機具(GV-3-)": "#D7BDE2", "鍋爐(GS-1-)": "#E6B0AA", "發電機(GS-2-)": "#A9CCE3", "肩背或手持式割草機、吹葉機(GS-3-)": "#A3E4D7", "肩背或手持式農用機具(GS-4-)": "#F9E79F" }
DASH_PALETTE = ['#B0C4DE', '#F5CBA7', '#A9CCE3', '#E6B0AA', '#D7BDE2', '#A3E4D7', '#F9E79F', '#95A5A6', '#85C1E9', '#D2B4DE', '#F1948A', '#76D7C4']
UNREPORTED_COLORS = ["#D5DBDB", "#FAD7A0", "#D2B4DE", "#AED6F1", "#A3E4D7", "#F5B7B1"]

@st.cache_resource
def init_google_fuel():
    oauth = st.secrets["gcp_oauth"]
    creds = Credentials(token=None, refresh_token=oauth["refresh_token"], token_uri="https://oauth2.googleapis.com/token", client_id=oauth["client_id"], client_secret=oauth["client_secret"], scopes=["https://www.googleapis.com/auth/drive", "https://www.googleapis.com/auth/spreadsheets"])
    gc = gspread.authorize(creds); drive = build('drive', 'v3', credentials=creds)
    return gc, drive

try:
    gc, drive_service = init_google_fuel()
    sh = gc.open_by_key(SHEET_ID)
    try: ws_equip = sh.worksheet("設備清單") 
    except: ws_equip = sh.sheet1 
    try: ws_record = sh.worksheet("油料填報紀錄")
    except: ws_record = sh.worksheet("填報紀錄")
except Exception as e: st.error(f"連線失敗: {e}"); st.stop()

@st.cache_data(ttl=600)
def load_fuel_data():
    df_e = pd.DataFrame(ws_equip.get_all_records()).astype(str)
    if '設備編號' in df_e.columns: df_e['統計類別'] = df_e['設備編號'].apply(lambda c: next((v for k, v in DEVICE_CODE_MAP.items() if str(c).startswith(k)), "其他/未分類"))
    df_e['設備數量_num'] = pd.to_numeric(df_e['設備數量'], errors='coerce').fillna(1)
    
    data = ws_record.get_all_values()
    df_r = pd.DataFrame(data[1:], columns=data[0]) if len(data) > 1 else pd.DataFrame(columns=data[0])
    return df_e, df_r

df_equip, df_records = load_fuel_data()

# ==========================================
# 4. Word 匯出輔助函數 (整合多頁 PDF 支援)
# ==========================================
def get_drive_id(url):
    match = re.search(r'/d/([a-zA-Z0-9_-]+)', str(url))
    if match: return match.group(1)
    match = re.search(r'id=([a-zA-Z0-9_-]+)', str(url))
    if match: return match.group(1)
    return None

def download_and_convert_drive_files(drive_service_obj, file_id):
    try:
        request = drive_service_obj.files().get_media(fileId=file_id)
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        done = False
        while done is False:
            status, done = downloader.next_chunk()
        
        fh.seek(0)
        header = fh.read(4)
        fh.seek(0)
        
        if header == b'%PDF':
            images = []
            try:
                doc = fitz.open(stream=fh.read(), filetype="pdf")
                for page_num in range(len(doc)):
                    page = doc.load_page(page_num)
                    pix = page.get_pixmap(dpi=150)
                    img_bytes = pix.tobytes("png")
                    images.append(io.BytesIO(img_bytes))
                return images
            except:
                return []
        return [fh]
    except:
        return []

def export_general_docx(df_year, df_eq, drive_srv):
    doc = Document()
    for section in doc.sections:
        section.page_width = Cm(21.0); section.page_height = Cm(29.7)
        section.left_margin = Cm(1.5); section.right_margin = Cm(1.5); section.top_margin = Cm(1.5); section.bottom_margin = Cm(1.5)

    df_gen = df_year[~df_year['備註'].astype(str).str.contains('批次申報', na=False)].copy()
    df_merged = pd.merge(df_gen, df_eq[['設備名稱備註', '設備編號']], on='設備名稱備註', how='left')
    df_merged['設備編號'] = df_merged['設備編號'].fillna('無編號')
    df_merged['佐證資料_clean'] = df_merged['佐證資料'].fillna('無').astype(str).str.strip()

    valid_df = df_merged[(df_merged['佐證資料_clean'] != '') & (df_merged['佐證資料_clean'] != '無')]
    groups = valid_df.groupby('佐證資料_clean')
    group_list = []
    
    for link_str, group in groups:
        eqs = group[['設備編號', '設備名稱備註', '填報單位', '原燃物料名稱', '與其他設備共用加油單', '備註']].drop_duplicates().sort_values('設備編號')
        group_list.append({'min_eq_id': eqs['設備編號'].min(), 'link_str': link_str, 'equipments': eqs})
        
    group_list.sort(key=lambda x: str(x['min_eq_id']))
    
    for g in group_list:
        p = doc.add_paragraph()
        for _, eq in g['equipments'].iterrows():
            yearly_vol = df_year[df_year['設備名稱備註'] == eq['設備名稱備註']]['加油量'].sum()
            run = p.add_run(f"填報單位：{eq['填報單位']} | 設備名稱：{eq['設備名稱備註']} ({eq['設備編號']})\n")
            run.bold = True
            p.add_run(f"燃料：{eq['原燃物料名稱']} | 年度總加油量：{yearly_vol:,.1f} 公升\n").bold = True
            if str(eq['與其他設備共用加油單']) == "是":
                p.add_run(f"⚠️ 此為共用加油單 (備註：{eq['備註']})\n").bold = True
            
        links = g['link_str'].split('\n')
        images = []
        for link in links:
            fid = get_drive_id(link)
            if fid:
                images.extend(download_and_convert_drive_files(drive_srv, fid))
                
        if len(images) == 1:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try: p_img.add_run().add_picture(images[0], height=Cm(11.0))
            except: pass
        elif len(images) > 1:
            table = doc.add_table(rows=0, cols=2)
            table.autofit = False
            for i, img in enumerate(images):
                if i % 2 == 0: row_cells = table.add_row().cells
                try:
                    p_img = row_cells[i % 2].paragraphs[0]
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.add_run().add_picture(img, height=Cm(7.5))
                except: pass
        doc.add_page_break()
        
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def export_batch_docx(df_year, drive_srv):
    doc = Document()
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Cm(29.7); section.page_height = Cm(21.0)
        section.left_margin = Cm(1.5); section.right_margin = Cm(1.5); section.top_margin = Cm(1.5); section.bottom_margin = Cm(1.5)

    df_batch = df_year[df_year['備註'].astype(str).str.contains('批次申報', na=False)].copy()
    if df_batch.empty: return None

    df_batch['批次類別'] = df_batch['備註'].apply(lambda x: str(x).split(' | ')[0] if ' | ' in str(x) else str(x))
    groups = df_batch.groupby(['填報單位', '原燃物料名稱', '批次類別'])

    for name, group in groups:
        dept, fuel, cat = name
        eq_names = group['設備名稱備註'].unique()
        eq_str = "、".join(eq_names)
        yearly_vol = group['加油量'].sum()
        
        p = doc.add_paragraph()
        p.add_run(f"填報單位：{dept}\n").bold = True
        p.add_run(f"設備名稱：{eq_str}\n").bold = True
        p.add_run(f"燃料：{fuel} | 年度總加油量：{yearly_vol:,.1f} 公升\n").bold = True
        
        unique_links = group['佐證資料'].dropna().unique()
        images = []
        for link in unique_links:
            if str(link).strip() in ["無", ""]: continue
            for l in str(link).split('\n'):
                fid = get_drive_id(l)
                if fid: images.extend(download_and_convert_drive_files(drive_srv, fid))
        
        if len(images) == 1:
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try: p_img.add_run().add_picture(images[0], width=Cm(22.0))
            except: pass
        elif len(images) > 1:
            table = doc.add_table(rows=0, cols=2)
            table.autofit = False
            for i, img in enumerate(images):
                if i % 2 == 0: row_cells = table.add_row().cells
                try:
                    p_img = row_cells[i % 2].paragraphs[0]
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.add_run().add_picture(img, width=Cm(12.0))
                except: pass
        doc.add_page_break()

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# ==========================================
# 5. 後台主程式
# ==========================================
def render_admin_dashboard():
    df_clean = df_records.copy()
    if not df_clean.empty:
        df_clean['加油量'] = pd.to_numeric(df_clean['加油量'], errors='coerce').fillna(0)
        df_clean['日期格式'] = pd.to_datetime(df_clean['加油日期'], errors='coerce')
        df_clean['年份'] = df_clean['日期格式'].dt.year.fillna(0).astype(int)
        df_clean['月份'] = df_clean['日期格式'].dt.month.fillna(0).astype(int)
        df_clean['油品大類'] = df_clean['原燃物料名稱'].apply(lambda x: '汽油' if '汽油' in str(x) else ('柴油' if '柴油' in str(x) else '其他'))
        if not df_equip.empty:
            device_map = pd.Series(df_equip['統計類別'].values, index=df_equip['設備名稱備註']).to_dict()
            df_clean['統計類別'] = df_clean['設備名稱備註'].map(device_map).fillna("其他/未分類")

    all_years = sorted(df_clean['年份'][df_clean['年份']>0].unique(), reverse=True) if not df_clean.empty else [datetime.now().year]
    c_year, _ = st.columns([1, 3])
    selected_admin_year = c_year.selectbox("📅 請選擇檢視年度", all_years, index=0)
    df_year = df_clean[df_clean['年份'] == selected_admin_year] if not df_clean.empty else pd.DataFrame()

    admin_tabs = st.tabs(["🔍 申報資料異動", "⚠️ 篩選未申報名單", "📝 全校燃油設備總覽", "📊 全校油料使用儀表板"])

    # === Tab 1: 申報資料異動 ===
    with admin_tabs[0]:
        st.subheader("🔍 申報資料異動")
        if not df_year.empty:
            df_year['加油日期'] = pd.to_datetime(df_year['加油日期']).dt.date
            edited = st.data_editor(df_year, column_config={"佐證資料": st.column_config.LinkColumn("佐證", display_text="🔗"), "加油日期": st.column_config.DateColumn("日期", format="YYYY-MM-DD"), "加油量": st.column_config.NumberColumn("油量", format="%.2f"), "填報時間": st.column_config.TextColumn("填報時間", disabled=True)}, num_rows="dynamic", use_container_width=True, key="editor_v122")
            
            c_btn1, c_btn2 = st.columns([2, 10])
            with c_btn1:
                df_stats = df_year.groupby(['設備名稱備註'])['加油量'].sum().reset_index()
                df_export = pd.merge(df_equip, df_stats, on='設備名稱備註', how='left')
                df_export['加油量'] = df_export['加油量'].fillna(0)
                df_export.rename(columns={'加油量': f'{selected_admin_year}年度總加油量'}, inplace=True)
                target_cols = ['填報單位', '設備名稱備註', '校內財產編號', '原燃物料名稱', '保管人', '設備所屬單位/部門', '設備詳細位置/樓層', '設備數量', '設備編號', f'{selected_admin_year}年度總加油量']
                df_final_export = df_export[[c for c in target_cols if c in df_export.columns]]
                csv_data = df_final_export.to_csv(index=False).encode('utf-8-sig')
                st.download_button(label="📥 下載設備年度加油統計表", data=csv_data, file_name=f"{selected_admin_year}_設備年度加油統計.csv", mime="text/csv")
            with c_btn2:
                if st.button("💾 儲存變更", type="primary"):
                    try:
                        df_all_data = df_records.copy()
                        df_all_data['temp_date'] = pd.to_datetime(df_all_data['加油日期'], errors='coerce')
                        df_all_data['temp_year'] = df_all_data['temp_date'].dt.year.fillna(0).astype(int)
                        df_keep = df_all_data[df_all_data['temp_year'] != selected_admin_year].copy()
                        df_new = edited.copy()
                        df_final = pd.concat([df_keep, df_new], ignore_index=True)
                        if 'temp_date' in df_final.columns: del df_final['temp_date']
                        if 'temp_year' in df_final.columns: del df_final['temp_year']
                        if '加油日期' in df_final.columns: df_final['加油日期'] = df_final['加油日期'].astype(str)
                        df_final = df_final[df_records.columns.tolist()].sort_values(by='加油日期', ascending=False)

                        ws_record.clear()
                        ws_record.update([df_final.columns.tolist()] + df_final.astype(str).values.tolist())
                        st.success("✅ 更新成功！資料已安全合併存檔。"); st.cache_data.clear(); time.sleep(1); st.rerun()
                    except Exception as e: st.error(f"更新失敗: {e}")
            
            st.markdown("---")
            st.markdown("##### 📁 佐證資料匯出")
            if not DOCX_READY:
                st.error("⚠️ 系統尚未安裝 `python-docx` 或 `pymupdf` 套件。請確認已將其加入 `requirements.txt` 並重新部署。")
            else:
                c_w1, c_w2 = st.columns(2)
                with c_w1:
                    if st.button("⚡ 產生【一般申報】佐證資料", use_container_width=True):
                        with st.spinner("正在下載圖片並合併 (包含自動轉檔 PDF)，可能需要幾分鐘..."):
                            st.session_state['doc_general'] = export_general_docx(df_year, df_equip, drive_service)
                    if 'doc_general' in st.session_state:
                        st.download_button("⬇️ 點擊下載【一般申報】Word", data=st.session_state['doc_general'], file_name=f"{selected_admin_year}_一般申報佐證資料.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)

                with c_w2:
                    if st.button("⚡ 產生【油卡批次申報】佐證資料", use_container_width=True):
                        with st.spinner("正在下載圖片並合併 (包含自動轉檔 PDF)，可能需要幾分鐘..."):
                            st.session_state['doc_batch'] = export_batch_docx(df_year, drive_service)
                    if 'doc_batch' in st.session_state:
                        st.download_button("⬇️ 點擊下載【油卡批次申報】Word", data=st.session_state['doc_batch'], file_name=f"{selected_admin_year}_油卡批次申報佐證資料.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)

        else: st.info(f"{selected_admin_year} 年度尚無資料。")

    # === Tab 2, 3, 4 保留原本的所有圖表與邏輯 (避免篇幅過長，這裡完美銜接上述程式邏輯) ===
    # ... Tab 2: 篩選未申報名單 ...
    with admin_tabs[1]:
        st.subheader("⚠️ 篩選未申報名單")
        c_f1, c_f2 = st.columns(2)
        d_start = c_f1.date_input("查詢起始日", date(selected_admin_year, 1, 1))
        d_end = c_f2.date_input("查詢結束日", date.today())
        if st.button("開始篩選"):
            if not df_clean.empty:
                mask = (df_clean['日期格式'].dt.date >= d_start) & (df_clean['日期格式'].dt.date <= d_end)
                reported = set(df_clean[mask]['設備名稱備註'].unique())
                df_eq_copy = df_equip.copy()
                df_eq_copy['已申報'] = df_eq_copy['設備名稱備註'].apply(lambda x: x in reported)
                unreported = df_eq_copy[~df_eq_copy['已申報']]
                if not unreported.empty:
                    st.error(f"🚩 期間 [{d_start} ~ {d_end}] 共有 {len(unreported)} 台設備未申報！")
                    for idx, (unit, group) in enumerate(unreported.groupby('填報單位')):
                        bg_color = UNREPORTED_COLORS[idx % len(UNREPORTED_COLORS)]
                        st.markdown(f"""<div class="unreported-block" style="background-color: {bg_color};"><div class="unreported-title">🏢 {unit} (未申報數: {len(group)})</div></div>""", unsafe_allow_html=True)
                        st.dataframe(group[['設備名稱備註', '保管人', '校內財產編號']], use_container_width=True)
                else: st.success("🎉 太棒了！全數已申報。")
            else: st.warning("無資料可供篩選。")

    # ... Tab 3: 全校燃油設備總覽 ...
    with admin_tabs[2]:
        if not df_year.empty and not df_equip.empty:
            total_eq = int(df_equip['設備數量_num'].sum())
            gas_eq = int(df_equip[df_equip['原燃物料名稱'].str.contains('汽油', na=False)]['設備數量_num'].sum())
            diesel_eq = int(df_equip[df_equip['原燃物料名稱'].str.contains('柴油', na=False)]['設備數量_num'].sum())
            k1, k2, k3 = st.columns(3)
            k1.markdown(f"""<div class="top-kpi-card"><div class="top-kpi-title">🚜 全校燃油設備總數</div><div class="top-kpi-value">{total_eq}</div></div>""", unsafe_allow_html=True)
            k2.markdown(f"""<div class="top-kpi-card"><div class="top-kpi-title">⛽ 全校汽油設備數</div><div class="top-kpi-value">{gas_eq}</div></div>""", unsafe_allow_html=True)
            k3.markdown(f"""<div class="top-kpi-card"><div class="top-kpi-title">🚛 全校柴油設備數</div><div class="top-kpi-value">{diesel_eq}</div></div>""", unsafe_allow_html=True)
            st.markdown("---")

            st.subheader("📂 各類設備數量及用油統計")
            eq_sums = df_equip.groupby('統計類別')['設備數量_num'].sum()
            eq_gas_sums = df_equip[df_equip['原燃物料名稱'].str.contains('汽油', na=False)].groupby('統計類別')['設備數量_num'].sum()
            eq_dsl_sums = df_equip[df_equip['原燃物料名稱'].str.contains('柴油', na=False)].groupby('統計類別')['設備數量_num'].sum()
            fuel_sums = df_year.groupby(['統計類別', '油品大類'])['加油量'].sum().unstack(fill_value=0)
            
            for i in range(0, len(DEVICE_ORDER), 2):
                cols = st.columns(2)
                for j in range(2):
                    if i + j < len(DEVICE_ORDER):
                        category = DEVICE_ORDER[i + j]
                        with cols[j]:
                            count_tot = int(eq_sums.get(category, 0))
                            count_gas = int(eq_gas_sums.get(category, 0))
                            count_dsl = int(eq_dsl_sums.get(category, 0))
                            gas_vol = fuel_sums.loc[category, '汽油'] if category in fuel_sums.index and '汽油' in fuel_sums.columns else 0
                            diesel_vol = fuel_sums.loc[category, '柴油'] if category in fuel_sums.index and '柴油' in fuel_sums.columns else 0
                            total_vol = gas_vol + diesel_vol
                            header_color = MORANDI_COLORS.get(category, "#CFD8DC")
                            st.markdown(f"""<div class="stat-card-v119"><div class="stat-header" style="background-color: {header_color};"><span class="stat-title">{category}</span><span class="stat-count">{count_tot}</span></div><div class="stat-body-split"><div class="stat-col-left"><div class="stat-item"><span class="stat-item-label">⛽ 汽油設備數</span><span class="stat-item-val">{count_gas}</span></div><div class="stat-item"><span class="stat-item-label">🚛 柴油設備數</span><span class="stat-item-val">{count_dsl}</span></div><div class="stat-item"><span class="stat-item-label">🔥 燃油設備數</span><span class="stat-item-val">{count_tot}</span></div></div><div class="stat-col-right"><div class="stat-item"><span class="stat-item-label">汽油加油量(公升)</span><span class="stat-item-val">{gas_vol:,.1f}</span></div><div class="stat-item"><span class="stat-item-label">柴油加油量(公升)</span><span class="stat-item-val">{diesel_vol:,.1f}</span></div><div class="stat-item"><span class="stat-item-label">總計加油量(公升)</span><span class="stat-item-val">{total_vol:,.1f}</span></div></div></div></div>""", unsafe_allow_html=True)
            
            st.markdown("---")
            st.subheader("📋 各類設備詳細申報紀錄一覽")
            
            for category in DEVICE_ORDER:
                target_devices = df_equip[df_equip['統計類別'] == category]
                if not target_devices.empty:
                    st.markdown(f"<h3 style='color: #2874A6; font-size: 1.6rem; margin-top: 40px; margin-bottom: 25px;'>{category}</h3>", unsafe_allow_html=True)
                    
                    device_list = []
                    for _, row in target_devices.iterrows():
                        d_name = row['設備名稱備註']
                        d_id = row.get('設備編號', '無編號')
                        d_unit = row.get('填報單位', '-')
                        d_sub = row.get('設備所屬單位/部門', '-')
                        d_keeper = row.get('保管人', '-')
                        d_qty = row.get('設備數量', '1')
                        raw_fuel = row.get('原燃物料名稱', '-')
                        d_fuel = '汽油' if '汽油' in raw_fuel else ('柴油' if '柴油' in raw_fuel else raw_fuel)
                        d_vol = df_year[df_year['設備名稱備註'] == d_name]['加油量'].sum()
                        d_count = len(df_year[df_year['設備名稱備註'] == d_name])
                        status_html = '<span class="alert-status">⚠️ 尚未申報</span>' if d_count == 0 else ""

                        device_list.append({ "id": d_id, "name": d_name, "vol": d_vol, "fuel": d_fuel, "unit": d_unit, "sub": d_sub, "keeper": d_keeper, "qty": d_qty, "count": d_count, "status": status_html })
                    
                    for k in range(0, len(device_list), 2):
                        d_cols = st.columns(2)
                        for m in range(2):
                            if k + m < len(device_list):
                                item = device_list[k + m]
                                with d_cols[m]:
                                    st.markdown(f"""<div class="dev-card-v148"><div class="dev-header" style="background-color: {MORANDI_COLORS.get(category, '#34495E')};"><div class="dev-header-left"><div class="dev-id">{item['id']}</div><div class="dev-name-row"><span class="dev-name">{item['name']}</span><span class="qty-badge">數量: {item['qty']}</span></div></div><div class="dev-header-right"><div class="dev-vol">{item['vol']:.1f}<span class="dev-unit">公升</span></div></div></div><div class="dev-body"><div class="dev-item"><span class="dev-label">燃料種類:</span><span class="dev-val">{item['fuel']}</span></div><div class="dev-item"><span class="dev-label">填報單位:</span><span class="dev-val">{item['unit']}</span></div><div class="dev-item"><span class="dev-label">所屬部門:</span><span class="dev-val">{item['sub']}</span></div><div class="dev-item"><span class="dev-label">保管人:</span><span class="dev-val">{item['keeper']}</span></div></div><div class="dev-footer"><div class="dev-count">年度申報次數: {item['count']} 次</div><div>{item['status']}</div></div></div>""", unsafe_allow_html=True)

            st.markdown("---")
            st.subheader("📊 油品設備用油量佔比分析 (水平長條圖)")
            color_map = { "公務車輛(GV-1-)": "#B0C4DE", "乘坐式割草機(GV-2-)": "#F5CBA7", "乘坐式農用機具(GV-3-)": "#D7BDE2", "鍋爐(GS-1-)": "#E6B0AA", "發電機(GS-2-)": "#A9CCE3", "肩背或手持式割草機、吹葉機(GS-3-)": "#A3E4D7", "肩背或手持式農用機具(GS-4-)": "#F9E79F" }
            c_bar1, c_bar2 = st.columns(2)
            with c_bar1:
                st.markdown('<div class="bar-chart-box">', unsafe_allow_html=True)
                gas_data = df_year[(df_year['油品大類'] == '汽油') & (df_year['統計類別'].isin(DEVICE_ORDER))].groupby('統計類別')['加油量'].sum().reset_index()
                if not gas_data.empty:
                    gas_data = gas_data[gas_data['加油量'] > 0].sort_values('加油量', ascending=True)
                    total_g = gas_data['加油量'].sum()
                    gas_data['Label'] = gas_data['加油量'].apply(lambda x: f"{x:,.1f} L ({(x/total_g)*100:.1f}%)")
                    fig_g = px.bar(gas_data, x='加油量', y='統計類別', orientation='h', title='⛽ 汽油設備用油量佔比', color='統計類別', color_discrete_map=color_map, text='Label')
                    fig_g.update_layout(height=450, showlegend=False, xaxis=dict(title="加油量 (公升)", showgrid=True, gridcolor='#EAEDED'), yaxis=dict(title=""), plot_bgcolor='rgba(0,0,0,0)', font=dict(size=14))
                    fig_g.update_traces(textposition='outside', textfont=dict(size=14, color='black'))
                    fig_g.update_xaxes(range=[0, gas_data['加油量'].max() * 1.35])
                    st.plotly_chart(fig_g, use_container_width=True)
                else: st.info("無汽油數據")
                st.markdown('</div>', unsafe_allow_html=True)
            with c_bar2:
                st.markdown('<div class="bar-chart-box">', unsafe_allow_html=True)
                dsl_data = df_year[(df_year['油品大類'] == '柴油') & (df_year['統計類別'].isin(DEVICE_ORDER))].groupby('統計類別')['加油量'].sum().reset_index()
                if not dsl_data.empty:
                    dsl_data = dsl_data[dsl_data['加油量'] > 0].sort_values('加油量', ascending=True)
                    total_d = dsl_data['加油量'].sum()
                    dsl_data['Label'] = dsl_data['加油量'].apply(lambda x: f"{x:,.1f} L ({(x/total_d)*100:.1f}%)")
                    fig_d = px.bar(dsl_data, x='加油量', y='統計類別', orientation='h', title='🚛 柴油設備用油量佔比', color='統計類別', color_discrete_map=color_map, text='Label')
                    fig_d.update_layout(height=450, showlegend=False, xaxis=dict(title="加油量 (公升)", showgrid=True, gridcolor='#EAEDED'), yaxis=dict(title=""), plot_bgcolor='rgba(0,0,0,0)', font=dict(size=14))
                    fig_d.update_traces(textposition='outside', textfont=dict(size=14, color='black'))
                    fig_d.update_xaxes(range=[0, dsl_data['加油量'].max() * 1.35])
                    st.plotly_chart(fig_d, use_container_width=True)
                else: st.info("無柴油數據")
                st.markdown('</div>', unsafe_allow_html=True)
        else: st.warning("尚無資料可供統計。")

    # ... Tab 4: 全校油料使用儀表板 ...
    with admin_tabs[3]:
        if not df_year.empty:
            st.markdown(f"<div class='dashboard-main-title'>{selected_admin_year}年度 能源使用與碳排統計</div>", unsafe_allow_html=True)
            gas_sum = df_year[df_year['油品大類'] == '汽油']['加油量'].sum()
            diesel_sum = df_year[df_year['油品大類'] == '柴油']['加油量'].sum()
            total_sum = df_year['加油量'].sum()
            total_co2 = (gas_sum * 0.0022) + (diesel_sum * 0.0027)
            gas_pct = (gas_sum / total_sum * 100) if total_sum > 0 else 0
            diesel_pct = (diesel_sum / total_sum * 100) if total_sum > 0 else 0
            
            c1, c2 = st.columns(2)
            with c1: st.markdown(f"""<div class="admin-kpi-card"><div class="admin-kpi-header" style="background-color: #B0C4DE;">⛽ 汽油使用量</div><div class="admin-kpi-body"><div class="admin-kpi-value">{gas_sum:,.2f}<span class="admin-kpi-unit">公升</span></div><div class="admin-kpi-sub">佔比 {gas_pct:.1f}%</div></div></div>""", unsafe_allow_html=True)
            with c2: st.markdown(f"""<div class="admin-kpi-card"><div class="admin-kpi-header" style="background-color: #F5CBA7;">🚛 柴油使用量</div><div class="admin-kpi-body"><div class="admin-kpi-value">{diesel_sum:,.2f}<span class="admin-kpi-unit">公升</span></div><div class="admin-kpi-sub">佔比 {diesel_pct:.1f}%</div></div></div>""", unsafe_allow_html=True)
            st.write("") 
            c3, c4 = st.columns(2)
            with c3: st.markdown(f"""<div class="admin-kpi-card"><div class="admin-kpi-header" style="background-color: #A9CCE3;">💧 總用油量</div><div class="admin-kpi-body"><div class="admin-kpi-value">{total_sum:,.2f}<span class="admin-kpi-unit">公升</span></div><div class="admin-kpi-sub">100%</div></div></div>""", unsafe_allow_html=True)
            with c4: st.markdown(f"""<div class="admin-kpi-card"><div class="admin-kpi-header" style="background-color: #E6B0AA;">☁️ 碳排放量</div><div class="admin-kpi-body"><div class="admin-kpi-value">{total_co2:,.4f}<span class="admin-kpi-unit">公噸CO<sub>2</sub>e</span></div><div class="admin-kpi-sub">ESG 指標</div></div></div>""", unsafe_allow_html=True)
            st.markdown("---")

            st.subheader("📈 全校逐月加油量統計")
            monthly = df_year.groupby(['月份', '油品大類'])['加油量'].sum().reset_index()
            full_months = pd.DataFrame({'月份': range(1, 13)})
            monthly = full_months.merge(monthly, on='月份', how='left').fillna({'加油量':0, '油品大類':'汽油'})
            fig_month = px.bar(monthly, x='月份', y='加油量', color='油品大類', barmode='group', text_auto='.2f', color_discrete_sequence=DASH_PALETTE)
            fig_month.update_layout(xaxis=dict(tickmode='linear', tick0=1, dtick=1, title_font=dict(size=20), tickfont=dict(size=18, color='#566573')), yaxis=dict(title="加油量(公升)", title_font=dict(size=20), tickfont=dict(size=18, color='#566573')), font=dict(size=18), showlegend=True, height=500, margin=dict(t=50))
            fig_month.update_traces(textfont=dict(color='black', size=14), textangle=0, textposition='outside')
            st.plotly_chart(fig_month, use_container_width=True)

            st.markdown("---")
            st.subheader("🏆 全校前十大加油量單位")
            top_fuel = st.radio("選擇油品類型", ["汽油", "柴油"], horizontal=True, label_visibility="collapsed")
            df_top = df_year[df_year['油品大類'] == top_fuel]
            if not df_top.empty:
                top10_data = df_top.groupby('填報單位')['加油量'].sum().nlargest(10).reset_index()
                fig_top = px.bar(top10_data, x='填報單位', y='加油量', text_auto='.2f', title=f"{top_fuel}用量前十大單位", color_discrete_sequence=DASH_PALETTE)
                fig_top.update_layout(xaxis=dict(categoryorder='total descending', title_font=dict(size=20), tickfont=dict(size=18, color='#566573')), yaxis=dict(title="加油量(公升)", title_font=dict(size=20), tickfont=dict(size=18, color='#566573')), font=dict(size=18), height=600, margin=dict(t=50))
                fig_top.update_traces(selector=dict(type='bar'), width=0.5, textposition='outside', textangle=0, textfont=dict(color='black', size=14))
                st.plotly_chart(fig_top, use_container_width=True)
            else: st.info("無此油品數據。")

            st.markdown("---")
            st.subheader("📊 全校加油量單位佔比 (水平長條圖)")
            c_bu1, c_bu2 = st.columns(2)
            with c_bu1:
                st.markdown('<div class="bar-chart-box">', unsafe_allow_html=True)
                df_gas = df_year[(df_year['油品大類'] == '汽油') & (df_year['加油量'] > 0)]
                if not df_gas.empty:
                    gas_u_data = df_gas.groupby('填報單位')['加油量'].sum().reset_index().sort_values('加油量', ascending=True)
                    total_gu = gas_u_data['加油量'].sum()
                    gas_u_data['Label'] = gas_u_data['加油量'].apply(lambda x: f"{x:,.1f} L ({(x/total_gu)*100:.1f}%)")
                    fig_dg = px.bar(gas_u_data, x='加油量', y='填報單位', orientation='h', title='⛽ 汽油用量分佈', color='填報單位', color_discrete_sequence=DASH_PALETTE, text='Label')
                    fig_dg.update_layout(height=550, showlegend=False, xaxis=dict(title="加油量 (公升)", showgrid=True, gridcolor='#EAEDED'), yaxis=dict(title=""), plot_bgcolor='rgba(0,0,0,0)', font=dict(size=14))
                    fig_dg.update_traces(textposition='outside', textfont=dict(size=14, color='black'))
                    fig_dg.update_xaxes(range=[0, gas_u_data['加油量'].max() * 1.35])
                    st.plotly_chart(fig_dg, use_container_width=True)
                else: st.info("無汽油數據")
                st.markdown('</div>', unsafe_allow_html=True)
                
            with c_bu2:
                st.markdown('<div class="bar-chart-box">', unsafe_allow_html=True)
                df_dsl = df_year[(df_year['油品大類'] == '柴油') & (df_year['加油量'] > 0)]
                if not df_dsl.empty:
                    dsl_u_data = df_dsl.groupby('填報單位')['加油量'].sum().reset_index().sort_values('加油量', ascending=True)
                    total_du = dsl_u_data['加油量'].sum()
                    dsl_u_data['Label'] = dsl_u_data['加油量'].apply(lambda x: f"{x:,.1f} L ({(x/total_du)*100:.1f}%)")
                    fig_dd = px.bar(dsl_u_data, x='加油量', y='填報單位', orientation='h', title='🚛 柴油用量分佈', color='填報單位', color_discrete_sequence=DASH_PALETTE, text='Label')
                    fig_dd.update_layout(height=550, showlegend=False, xaxis=dict(title="加油量 (公升)", showgrid=True, gridcolor='#EAEDED'), yaxis=dict(title=""), plot_bgcolor='rgba(0,0,0,0)', font=dict(size=14))
                    fig_dd.update_traces(textposition='outside', textfont=dict(size=14, color='black'))
                    fig_dd.update_xaxes(range=[0, dsl_u_data['加油量'].max() * 1.35])
                    st.plotly_chart(fig_dd, use_container_width=True)
                else: st.info("無柴油數據")
                st.markdown('</div>', unsafe_allow_html=True)

            st.markdown("---")
            st.subheader("🌍 全校油料使用碳排放量(公噸二氧化碳當量)結構")
            df_year['CO2e'] = df_year.apply(lambda r: r['加油量']*0.0022 if '汽油' in str(r['原燃物料名稱']) else r['加油量']*0.0027, axis=1)
            if not df_year.empty:
                fig_tree = px.treemap(df_year, path=['填報單位', '設備名稱備註'], values='CO2e', color='填報單位', color_discrete_sequence=DASH_PALETTE)
                fig_tree.update_traces(texttemplate='%{label}<br>%{value:.4f}<br>%{percentRoot:.1%}', textfont=dict(size=24))
                fig_tree.update_layout(height=700)
                st.plotly_chart(fig_tree, use_container_width=True)
            else: st.info("無數據")
        else: st.info("尚無該年度資料，無法顯示儀表板。")

if __name__ == "__main__":
    render_admin_dashboard()